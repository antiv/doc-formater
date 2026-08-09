"""Sintetički .docx dokumenti za testove.

Stvarni uzorak (vidi `SAMPLE_DOCX`) nema nijednu tabelu ni
fusnotu, pa bi invarijanta sadržaja na njemu bila prazna tvrdnja za te
elemente. Ovi fixture-i pokrivaju upravo one slučajeve na kojima naivno
brisanje "praznih" pasusa uništava dokument.
"""

from __future__ import annotations

import io
import zlib

import docx
from docx.oxml.ns import qn
from docx.shared import Cm

def _png_1px() -> bytes:
    """Najmanji validan PNG (1×1), da fixture ne zavisi od spoljnog fajla."""

    def chunk(tag: bytes, payload: bytes) -> bytes:
        body = tag + payload
        return (
            len(payload).to_bytes(4, "big")
            + body
            + (zlib.crc32(body) & 0xFFFFFFFF).to_bytes(4, "big")
        )

    ihdr = chunk(b"IHDR", (1).to_bytes(4, "big") + (1).to_bytes(4, "big") + bytes([8, 2, 0, 0, 0]))
    raw = b"\x00\xff\xff\xff"  # filter byte + jedan beli piksel
    idat = chunk(b"IDAT", zlib.compress(raw))
    iend = chunk(b"IEND", b"")
    return b"\x89PNG\r\n\x1a\n" + ihdr + idat + iend


def build_rich_document() -> bytes:
    """Dokument sa svim vrstama sadržaja koje brisanje sme da ugrozi."""
    document = docx.Document()

    document.add_paragraph("ALMA MATER EUROPAEA")
    document.add_paragraph("")  # prazan pasus na naslovnici — ne sme da nestane

    document.add_paragraph("ZAHVALNICA", style="Heading 1")
    document.add_paragraph("Hvala svima.")

    document.add_paragraph("1 UVOD", style="Heading 1")
    document.add_paragraph("Prvi pasus glavnog teksta.")
    document.add_paragraph("")  # prazan pasus u telu — sme da nestane
    document.add_paragraph("Drugi pasus glavnog teksta.")

    # Pasus koji `paragraph.text` prijavljuje kao prazan, a nosi sliku.
    image_paragraph = document.add_paragraph()
    image_paragraph.add_run().add_picture(io.BytesIO(_png_1px()), width=Cm(2))
    document.add_paragraph("Slika 1: Naslov slike")
    document.add_paragraph("Vir: Anketni upitnik")

    document.add_paragraph("2 REZULTATI", style="Heading 1")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Zaglavlje A"
    table.cell(0, 1).text = "Zaglavlje B"
    table.cell(1, 0).text = "Vrednost 1"
    table.cell(1, 1).text = "Vrednost 2"
    # Prazan pasus unutar ćelije — nikad se ne dira.
    table.cell(1, 0).add_paragraph("")

    document.add_paragraph("")

    # Pasus koji je prazan ali nosi ručni prelom strane.
    break_paragraph = document.add_paragraph()
    run = break_paragraph.add_run()
    br = run._element.makeelement(qn("w:br"), {qn("w:type"): "page"})
    run._element.append(br)

    document.add_paragraph("LITERATURA", style="Heading 1")
    entry = document.add_paragraph("Smith-Autard, J. (2002). ")
    entry.add_run("https://doi.org/10.1080/14647893")

    document.add_paragraph("PRILOGE", style="Heading 1")
    document.add_paragraph("")  # prazan pasus u prilozima — ne sme da nestane
    document.add_paragraph("1. Koliko dugo predajete?")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_numbered_headings_document() -> bytes:
    """Dokument bez ijednog Word heading stila -- samo ručna numeracija."""
    document = docx.Document()
    for text in [
        "ALMA MATER EUROPAEA",
        "ZAHVALNICA",
        "1 UVOD",
        "Uvodni pasus.",
        "1.1 Predmet rada",
        "Pasus o predmetu.",
        "2.1.3 Devetnaesti vek",
        "Pasus o veku.",
        "LITERATURA",
        "Smith-Autard, J. (2002).",
        "PRILOGE",
        "1. Koliko dugo predajete?",
        "2. Da li je to dovoljno?",
    ]:
        document.add_paragraph(text)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
