"""Testovi nad sintetičkim dokumentima.

Stvarni uzorak ne pogađa sve opasne slučajeve -- nema tabela, nema pasusa sa
slikom usred tela, nema praznog pasusa u ćeliji. Ovde se upravo ti slučajevi
proveravaju, jer na njima naivno brisanje "praznih" pasusa uništava dokument.
"""

from __future__ import annotations

import io
import unittest

import docx

from docformat.analyze.structure import Role, Section, analyze
from docformat.formatting.engine import format_document
from docformat.formatting.ops.cleanup import is_safe_to_delete
from docformat.rules import load_rule_set

from .fixtures import build_numbered_headings_document, build_rich_document
from .helpers import AMEU_PRESET, content_census, empty_paragraph_count, table_text, text_sequence


class RichDocumentTest(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        cls.rule_set = load_rule_set(AMEU_PRESET)
        cls.source = build_rich_document()
        cls.result = format_document(io.BytesIO(cls.source), cls.rule_set)
        cls.output = cls.result.to_bytes()

    def test_table_survives(self) -> None:
        before, after = content_census(self.source), content_census(self.output)
        self.assertEqual(1, before["tabele (w:tbl)"])
        self.assertEqual(before["tabele (w:tbl)"], after["tabele (w:tbl)"])
        self.assertEqual(table_text(self.source), table_text(self.output))

    def test_image_survives(self) -> None:
        before, after = content_census(self.source), content_census(self.output)
        self.assertEqual(1, before["slike (w:drawing)"])
        self.assertEqual(before["slike (w:drawing)"], after["slike (w:drawing)"])
        self.assertEqual(before["image parts"], after["image parts"])

    def test_paragraph_holding_the_image_is_not_deleted(self) -> None:
        """`paragraph.text` ga prijavljuje kao prazan — brisanje bi ubilo sliku."""
        document = docx.Document(io.BytesIO(self.source))
        infos = analyze(document, self.rule_set.rules)
        image_infos = [i for i in infos if i.has_content and not i.text]
        self.assertTrue(image_infos, "fixture nema pasus sa slikom bez teksta")
        for info in image_infos:
            safe, reason = is_safe_to_delete(info)
            self.assertFalse(safe, "pasus sa slikom je označen kao bezbedan za brisanje")
            self.assertIn("ugrađeni sadržaj", reason)

    def test_text_is_identical(self) -> None:
        self.assertEqual(text_sequence(self.source), text_sequence(self.output))

    def test_empty_paragraphs_outside_body_are_kept(self) -> None:
        """Naslovnica i prilozi drže raspored praznim redovima."""
        deleted_sections = {c.section for c in self.result.report.deletions}
        self.assertNotIn(Section.COVER.value, deleted_sections)
        self.assertNotIn(Section.APPENDIX.value, deleted_sections)
        self.assertTrue(self.result.report.deletions, "ništa nije obrisano u telu")

    def test_empty_paragraph_inside_table_cell_is_kept(self) -> None:
        document = docx.Document(io.BytesIO(self.output))
        cell_paragraphs = [
            p for t in document.tables for r in t.rows for c in r.cells for p in c.paragraphs
        ]
        self.assertTrue(
            any(not p.text.strip() for p in cell_paragraphs),
            "prazan pasus u ćeliji je obrisan",
        )

    def test_section_properties_paragraph_is_never_deleted(self) -> None:
        before, after = content_census(self.source), content_census(self.output)
        self.assertEqual(before["prelomi sekcija (w:sectPr)"], after["prelomi sekcija (w:sectPr)"])

    def test_body_empty_paragraphs_are_removed(self) -> None:
        self.assertLess(empty_paragraph_count(self.output), empty_paragraph_count(self.source))


class MissingSectionTest(unittest.TestCase):
    """Pravilnik imenuje sekciju drugačije nego dokument.

    Nastaje redovno: pravilnik propisuje „PRILOGE", autor napiše „PRILOG A".
    Prilozi tada ostaju pripisani literaturi, pa bi se brisanje praznih redova
    proširilo na obrasce i anketne upitnike koji njima drže raspored.
    """

    def setUp(self) -> None:
        self.rule_set = load_rule_set(AMEU_PRESET)
        keywords = self.rule_set.rules.structure_profile.section_keywords
        keywords.appendix = ["PRILOGE"]  # fixture koristi "PRILOGE" -> nađeno
        self.source = build_rich_document()

    def test_found_appendix_allows_bibliography_cleanup(self) -> None:
        result = format_document(io.BytesIO(self.source), self.rule_set)
        self.assertEqual([], [w for w in result.report.warnings if "APPENDIX" in w])

    def test_missing_appendix_is_reported(self) -> None:
        self.rule_set.rules.structure_profile.section_keywords.appendix = ["DODATKI"]
        result = format_document(io.BytesIO(self.source), self.rule_set)
        self.assertTrue(
            any("APPENDIX nije pronađena" in w for w in result.report.warnings),
            "nenađena sekcija nije prijavljena",
        )

    def test_missing_appendix_restricts_deletion_to_body(self) -> None:
        self.rule_set.rules.structure_profile.section_keywords.appendix = ["DODATKI"]
        result = format_document(io.BytesIO(self.source), self.rule_set)
        sections = {c.section for c in result.report.deletions}
        self.assertNotIn(
            Section.BIBLIOGRAPHY.value,
            sections,
            "brisanje se proširilo na sekciju koja verovatno sadrži priloge",
        )


class NumberedHeadingsTest(unittest.TestCase):
    """Dokument bez Word heading stilova mora dati istu klasifikaciju sekcija."""

    def setUp(self) -> None:
        self.rules = load_rule_set(AMEU_PRESET).rules
        self.document = docx.Document(io.BytesIO(build_numbered_headings_document()))
        self.infos = analyze(self.document, self.rules)

    def test_sections_are_detected(self) -> None:
        by_text = {i.text: i.section for i in self.infos if i.text}
        self.assertEqual(Section.FRONT_MATTER, by_text["ZAHVALNICA"])
        self.assertEqual(Section.BODY, by_text["1 UVOD"])
        self.assertEqual(Section.BODY, by_text["Uvodni pasus."])
        self.assertEqual(Section.BIBLIOGRAPHY, by_text["LITERATURA"])
        self.assertEqual(Section.APPENDIX, by_text["PRILOGE"])

    def test_numbered_headings_are_detected(self) -> None:
        levels = {i.text: i.heading_level for i in self.infos if i.role is Role.HEADING}
        self.assertEqual(1, levels["1 UVOD"])
        self.assertEqual(2, levels["1.1 Predmet rada"])
        self.assertEqual(3, levels["2.1.3 Devetnaesti vek"])

    def test_numbered_questions_are_not_headings(self) -> None:
        """Anketna pitanja počinju brojem, ali nisu naslovi."""
        roles = {i.text: i.role for i in self.infos if i.text}
        self.assertEqual(Role.BODY_TEXT, roles["1. Koliko dugo predajete?"])
        self.assertEqual(Role.BODY_TEXT, roles["2. Da li je to dovoljno?"])


if __name__ == "__main__":
    unittest.main()
