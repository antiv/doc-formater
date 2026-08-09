#!/usr/bin/env python3
"""Poređenje novog formatiranja sa rezultatom stare skripte.

    .venv/bin/python tests/compare_with_legacy.py

Nije unittest nego dijagnostički alat: stara skripta je menjala tekst, pa
potpuna parnost nije ni cilj. Cilj je da se svako odstupanje objasni --
ili je namerno (stara skripta je prepisivala tekst), ili je regresija.
"""

from __future__ import annotations

import os
import sys
from collections import Counter
from pathlib import Path

import docx
from docx.shared import Length

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from styleguard.formatting.engine import format_document  # noqa: E402
from styleguard.rules import load_rule_set  # noqa: E402

ROOT = Path(__file__).resolve().parent.parent
# Dokumenti nisu deo repozitorijuma (lični su podaci); putanje se po potrebi
# zadaju kroz okruženje.
SOURCE = ROOT / os.getenv("SAMPLE_DOCX", "sample.docx")
LEGACY = ROOT / os.getenv("LEGACY_DOCX", "sample-legacy.docx")
PRESET = ROOT / "presets" / "ameu.json"


def _pt(value):
    return round(value.pt, 2) if isinstance(value, Length) else value


def paragraph_profile(paragraph):
    """Stilski otisak pasusa -- ono što poređenje treba da uporedi."""
    fmt = paragraph.paragraph_format
    runs = paragraph.runs
    first = runs[0] if runs else None
    return {
        "alignment": int(fmt.alignment) if fmt.alignment is not None else None,
        "line_spacing": _pt(fmt.line_spacing),
        "space_before": _pt(fmt.space_before),
        "space_after": _pt(fmt.space_after),
        "font_name": first.font.name if first else None,
        "font_size": _pt(first.font.size) if first else None,
        "bold": first.font.bold if first else None,
        "italic": first.font.italic if first else None,
    }


def indexed_paragraphs(path_or_bytes) -> dict[str, list[dict]]:
    """Profili grupisani po tekstu pasusa, da se porede isti pasusi."""
    document = (
        docx.Document(path_or_bytes)
        if not isinstance(path_or_bytes, (str, Path))
        else docx.Document(str(path_or_bytes))
    )
    out: dict[str, list[dict]] = {}
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            out.setdefault(text, []).append(paragraph_profile(paragraph))
    return out


def main() -> int:
    if not SOURCE.exists() or not LEGACY.exists():
        print(
            "Nema uzorka za poređenje. Ovaj alat traži originalni .docx i izlaz "
            "stare skripte; dokumenti nisu deo repozitorijuma.\n"
            f"  SAMPLE_DOCX={SOURCE.name} (postoji: {SOURCE.exists()})\n"
            f"  LEGACY_DOCX={LEGACY.name} (postoji: {LEGACY.exists()})"
        )
        return 1

    rule_set = load_rule_set(PRESET)
    result = format_document(SOURCE, rule_set)

    legacy = indexed_paragraphs(LEGACY)
    import io

    new = indexed_paragraphs(io.BytesIO(result.to_bytes()))

    only_legacy = set(legacy) - set(new)
    only_new = set(new) - set(legacy)
    shared = set(legacy) & set(new)

    print(f"Pasusa po tekstu: stara={len(legacy)} nova={len(new)} zajedničkih={len(shared)}")
    print()

    if only_legacy:
        print(f"Samo u staroj verziji ({len(only_legacy)}) — očekivano tamo gde je")
        print("stara skripta prepisivala tekst (UPPERCASE naslovi, Grafikon→Slika,")
        print("preimenovanje PRILOG A/B, crtica→pomišljaj):")
        for text in sorted(only_legacy)[:12]:
            print(f"    {text[:90]!r}")
        if len(only_legacy) > 12:
            print(f"    … i još {len(only_legacy) - 12}")
        print()

    if only_new:
        print(f"Samo u novoj verziji ({len(only_new)}) — pasusi koje je stara skripta")
        print("obrisala (PRILOG C/D) ili prepisala:")
        for text in sorted(only_new)[:12]:
            print(f"    {text[:90]!r}")
        if len(only_new) > 12:
            print(f"    … i još {len(only_new) - 12}")
        print()

    differences: Counter[str] = Counter()
    examples: dict[str, tuple[str, object, object]] = {}
    for text in shared:
        legacy_profile, new_profile = legacy[text][0], new[text][0]
        for key in legacy_profile:
            if legacy_profile[key] != new_profile[key]:
                differences[key] += 1
                examples.setdefault(key, (text, legacy_profile[key], new_profile[key]))

    print(f"Stilska odstupanja na {len(shared)} zajedničkih pasusa:")
    if not differences:
        print("    nema — potpuna parnost")
    for key, count in differences.most_common():
        text, old, fresh = examples[key]
        share = 100 * count / len(shared)
        print(f"    {key:14} {count:5} ({share:5.1f}%)   stara={old!r} nova={fresh!r}")
        print(f"                   npr. {text[:70]!r}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
