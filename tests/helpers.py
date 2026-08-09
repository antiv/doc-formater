"""Pomoćne funkcije za testove -- merenje invarijanti nad .docx fajlom."""

from __future__ import annotations

import io
import os
import unittest
from pathlib import Path

import docx
from docx.oxml.ns import qn

PROJECT_ROOT = Path(__file__).resolve().parent.parent
AMEU_PRESET = PROJECT_ROOT / "presets" / "ameu.json"

# Uzorak je stvaran diplomski rad i namerno nije u repozitorijumu -- lični je
# podatak. Testovi koji ga traže preskaču se na svežem klonu; sve invarijante
# koje bi on pokrivao pokrivene su i sintetičkim fixture-om u `test_fixtures`.
SAMPLE_DOCX = PROJECT_ROOT / os.getenv(
    "SAMPLE_DOCX", "sample.docx"
)
SAMPLE_AVAILABLE = SAMPLE_DOCX.exists()
SKIP_REASON = (
    f"nema uzorka {SAMPLE_DOCX.name} (dokumenti nisu deo repozitorijuma; "
    "postavi SAMPLE_DOCX na svoj .docx da bi se ovi testovi izvršili)"
)


def requires_sample(test_item):
    """Dekorator: preskoči test kad uzorak nije dostupan."""
    return unittest.skipUnless(SAMPLE_AVAILABLE, SKIP_REASON)(test_item)

# Elementi čiji broj mora ostati nepromenjen: slike, grafikoni, tabele,
# fusnote, linkovi. Ovo je merenje zahteva "slike, tabele i grafike ne
# izbacivati iz fajla".
_CONTENT_ELEMENTS = {
    "slike (w:drawing)": "w:drawing",
    "VML slike (w:pict)": "w:pict",
    "objekti (w:object)": "w:object",
    "tabele (w:tbl)": "w:tbl",
    "fusnote": "w:footnoteReference",
    "endnote": "w:endnoteReference",
    "hyperlink-ovi": "w:hyperlink",
    "obeleživači": "w:bookmarkStart",
    "prelomi sekcija (w:sectPr)": "w:sectPr",
}


def open_document(source) -> docx.Document:
    if isinstance(source, (bytes, bytearray)):
        return docx.Document(io.BytesIO(source))
    return docx.Document(str(source))


def content_census(source) -> dict[str, int]:
    """Broj svakog nosećeg elementa u dokumentu."""
    document = open_document(source)
    body = document.element.body
    census = {
        label: len(body.findall(f".//{qn(tag)}"))
        for label, tag in _CONTENT_ELEMENTS.items()
    }
    # Slike se broje i preko relacija paketa -- `w:drawing` se može izgubiti a
    # da relacija ostane, i obrnuto.
    census["image parts"] = sum(
        1 for part in document.part.package.parts if part.content_type.startswith("image/")
    )
    return census


def text_sequence(source) -> list[str]:
    """Sekvenca nepraznih pasusa tela dokumenta, sa normalizovanim razmacima.

    Prazni pasusi se izostavljaju jer je njihovo brisanje dozvoljeno; sve
    ostalo mora ostati identično karakter po karakter.
    """
    document = open_document(source)
    out: list[str] = []
    for child in document.element.body.iterchildren():
        if child.tag != qn("w:p"):
            continue
        text = "".join(node.text or "" for node in child.iter(qn("w:t"))).strip()
        if text:
            out.append(text)
    return out


def table_text(source) -> list[str]:
    """Tekst svih ćelija svih tabela -- tabele se nikad ne diraju."""
    document = open_document(source)
    out: list[str] = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                out.append(cell.text.strip())
    return out


def empty_paragraph_count(source) -> int:
    document = open_document(source)
    count = 0
    for child in document.element.body.iterchildren():
        if child.tag != qn("w:p"):
            continue
        text = "".join(node.text or "" for node in child.iter(qn("w:t"))).strip()
        if not text:
            count += 1
    return count
