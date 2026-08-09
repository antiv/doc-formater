"""Formatiranje tabela.

Tabele se samo formatiraju -- nikad ne brišu, ne spajaju i ne prerađuju.
Obrađuju se i ugnežđene tabele, jer python-docx `document.tables` vraća samo
tabele najvišeg nivoa.
"""

from __future__ import annotations

from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table

from ...report import Change
from ...rules import FormattingRules
from ..runs import apply_paragraph_format, apply_run_format, iter_runs

_WD_TABLE_ALIGNMENT = {"LEFT": 0, "CENTER": 1, "RIGHT": 2}


def _iter_tables(container, document: DocxDocument):
    """Sve tabele, uključujući ugnežđene."""
    for table in container.tables:
        yield table
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_tables(cell, document)


def apply(document: DocxDocument, infos, rules: FormattingRules) -> list[Change]:
    config = rules.tables
    changes: list[Change] = []

    for table_index, table in enumerate(_iter_tables(document, document)):
        ctx_table = {"paragraph_index": None, "section": f"tabela {table_index + 1}", "role": "TABLE"}

        if config.alignment is not None:
            target = _WD_TABLE_ALIGNMENT.get(config.alignment.value)
            current = int(table.alignment) if table.alignment is not None else None
            if target is not None and current != target:
                table.alignment = target
                changes.append(
                    Change(kind="style", rule_path="tables.alignment",
                           before=current, after=config.alignment.value, **ctx_table)
                )

        if config.header_row_repeat is True and table.rows:
            if _set_repeat_header(table.rows[0]):
                changes.append(
                    Change(kind="style", rule_path="tables.header_row_repeat",
                           before=False, after=True, **ctx_table)
                )

        for row_index, row in enumerate(table.rows):
            is_header = row_index == 0
            bold = config.header_row_bold if is_header else None

            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    ctx = {
                        "paragraph_index": None,
                        "section": f"tabela {table_index + 1}",
                        "role": "TABLE_HEADER" if is_header else "TABLE_CELL",
                    }
                    changes.extend(
                        apply_paragraph_format(
                            paragraph,
                            line_spacing=config.cell_line_spacing,
                            space_before_pt=config.cell_space_before_pt,
                            space_after_pt=config.cell_space_after_pt,
                            context=ctx,
                        )
                    )
                    for run in iter_runs(paragraph):
                        changes.extend(
                            apply_run_format(
                                run,
                                font_name=rules.typography.font_family,
                                size_pt=config.cell_size_pt,
                                bold=bold,
                                context=ctx,
                            )
                        )

    return changes


def _set_repeat_header(row) -> bool:
    """`w:tblHeader` -- ponavljanje reda zaglavlja na svakoj strani."""
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is not None:
        return False
    tr_pr.append(tr_pr.makeelement(qn("w:tblHeader"), {}))
    return True
