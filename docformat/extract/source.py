"""Ekstrakcija teksta iz pravilnika (PDF ili DOCX).

Jedan izlaz za oba formata (`RulesDocument`), koji koriste i Mate agent i
regex fallback. Mate `/v1/chat/completions` je text-only ruta -- ne-tekstualni
delovi poruke se odbacuju na serveru -- pa se PDF nikad ne šalje kao fajl nego
se ovde pretvara u tekst.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterator, Literal

BlockKind = Literal["paragraph", "table_row", "heading"]

# Ispod ovoga smatramo da PDF nema tekstualni sloj (skeniran dokument).
MIN_MEANINGFUL_CHARS = 200


@dataclass
class TextBlock:
    text: str
    page: int | None = None
    kind: BlockKind = "paragraph"


class NoTextLayerError(RuntimeError):
    """PDF je skeniran -- nema šta da se pošalje agentu ni da se regexuje."""


class RulesDocument:
    def __init__(self, blocks: list[TextBlock], filename: str) -> None:
        self.blocks = blocks
        self.filename = filename

    def __len__(self) -> int:
        return len(self.blocks)

    @property
    def char_count(self) -> int:
        return sum(len(b.text) for b in self.blocks)

    @property
    def plain_text(self) -> str:
        """Sirov tekst bez oznaka strana -- za verifikaciju citata i regex."""
        return "\n".join(b.text for b in self.blocks)

    def head(self, pages: int = 2, max_blocks: int = 60) -> "RulesDocument":
        """Početak dokumenta -- naslovna strana nosi naziv institucije.

        Za DOCX (gde broj strane ne postoji) pada na prvih `max_blocks`.
        """
        if any(b.page is not None for b in self.blocks):
            selected = [b for b in self.blocks if b.page is not None and b.page <= pages]
        else:
            selected = self.blocks[:max_blocks]
        return RulesDocument(selected, self.filename)

    def as_prompt_text(self, page_markers: bool = True) -> str:
        """Tekst za prompt, sa oznakama strana da agent može da citira stranu."""
        lines: list[str] = []
        current_page: int | None = object()  # sentinel
        for block in self.blocks:
            if page_markers and block.page != current_page:
                current_page = block.page
                if block.page is not None:
                    lines.append(f"\n=== STRANA {block.page} ===")
            prefix = "| " if block.kind == "table_row" else ""
            lines.append(f"{prefix}{block.text}")
        return "\n".join(lines).strip()


# --------------------------------------------------------------------------
# PDF
# --------------------------------------------------------------------------


def _serialize_table(table: list[list[str | None]]) -> Iterator[str]:
    """Tabelu u `Zaglavlje: vrednost` redove.

    Pravilnici tipografiju često drže u tabeli. Sirov `extract_text` iz nje
    vrati izmešane brojeve bez konteksta, pa se zaglavlje eksplicitno lepi uz
    svaku ćeliju.
    """
    if not table:
        return
    header = [(c or "").strip() for c in table[0]]
    has_header = sum(1 for c in header if c) >= 2

    for row in table[1:] if has_header else table:
        cells = [(c or "").strip().replace("\n", " ") for c in row]
        if not any(cells):
            continue
        if has_header:
            pairs = [
                f"{header[i]}: {cell}" if i < len(header) and header[i] else cell
                for i, cell in enumerate(cells)
                if cell
            ]
            yield " | ".join(pairs)
        else:
            yield " | ".join(c for c in cells if c)


def read_pdf(path: str | Path) -> RulesDocument:
    import pdfplumber

    path = Path(path)
    blocks: list[TextBlock] = []

    with pdfplumber.open(path) as pdf:
        for page_no, page in enumerate(pdf.pages, start=1):
            for table in page.extract_tables() or []:
                for row_text in _serialize_table(table):
                    blocks.append(TextBlock(row_text, page=page_no, kind="table_row"))

            text = page.extract_text() or ""
            for line in text.splitlines():
                line = line.strip()
                if line:
                    blocks.append(TextBlock(line, page=page_no, kind="paragraph"))

    doc = RulesDocument(blocks, path.name)
    if doc.char_count < MIN_MEANINGFUL_CHARS:
        raise NoTextLayerError(
            f"'{path.name}' nema tekstualni sloj (pronađeno {doc.char_count} karaktera). "
            "Verovatno je skeniran. Izaberi set iz biblioteke ili unesi pravila ručno."
        )
    return doc


# --------------------------------------------------------------------------
# DOCX
# --------------------------------------------------------------------------


def read_docx(path: str | Path) -> RulesDocument:
    """DOCX u blokove, uz očuvan redosled pasusa i tabela.

    Namerno se ne koristi `doc.paragraphs` -- taj property preskače sav
    sadržaj unutar tabela, a pravilnici u DOCX-u tipično imaju tipografiju
    upravo u tabeli.
    """
    import docx
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    path = Path(path)
    document = docx.Document(str(path))
    body = document.element.body
    blocks: list[TextBlock] = []

    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            paragraph = Paragraph(child, document)
            text = paragraph.text.strip()
            if not text:
                continue
            style_id = (paragraph.style.style_id or "") if paragraph.style else ""
            kind: BlockKind = "heading" if style_id.startswith("Heading") else "paragraph"
            blocks.append(TextBlock(text, page=None, kind=kind))

        elif child.tag == qn("w:tbl"):
            table = Table(child, document)
            grid = [[cell.text for cell in row.cells] for row in table.rows]
            for row_text in _serialize_table(grid):
                blocks.append(TextBlock(row_text, page=None, kind="table_row"))

    doc = RulesDocument(blocks, path.name)
    if doc.char_count < MIN_MEANINGFUL_CHARS:
        raise NoTextLayerError(f"'{path.name}' je prazan ili nečitljiv.")
    return doc


def read_rules_document(path: str | Path) -> RulesDocument:
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return read_pdf(path)
    if suffix == ".docx":
        return read_docx(path)
    raise ValueError(f"Nepodržan format pravilnika: '{suffix}'. Podržani su .pdf i .docx.")
