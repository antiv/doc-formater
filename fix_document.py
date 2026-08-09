import os
import re
import json
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def load_rules(json_path="ameu_rules.json"):
    if os.path.exists(json_path):
        with open(json_path, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}

def fix_page_ranges_in_text(text: str) -> str:
    """Zamenjuje obične crtice u rasponima stranica ili godina dugim stičnim pomišljajem (–)."""
    pattern = r'(\b\d+)\s*-\s*(\d+\b)'
    return re.sub(pattern, r'\1–\2', text)

def enable_mirror_margins(section):
    """Aktivira mirror margins u XML-u sekcije."""
    sectPr = section._sectPr
    if sectPr.find(qn('w:mirrorMargins')) is None:
        sectPr.append(OxmlElement('w:mirrorMargins'))

def remove_paragraph(paragraph):
    """Fizički uklanja odstavak iz XML stabla dokumenta."""
    p = paragraph._p
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)

def format_academic_paper(docx_input_path: str, docx_output_path: str, rules_path: str = "ameu_rules.json"):
    rules = load_rules(rules_path)
    doc = docx.Document(docx_input_path)
    font_name = rules.get("typography", {}).get("font_family", "Times New Roman")

    # 1. Postavka sekcija i zrcalnih robova (Mirror Margins: inside 3.0, top/bottom/outside 2.5)
    for section in doc.sections:
        section.top_margin = Cm(rules.get("page_setup", {}).get("margins_cm", {}).get("top", 2.5))
        section.bottom_margin = Cm(rules.get("page_setup", {}).get("margins_cm", {}).get("bottom", 2.5))
        section.left_margin = Cm(rules.get("page_setup", {}).get("margins_cm", {}).get("inside", 3.0))
        section.right_margin = Cm(rules.get("page_setup", {}).get("margins_cm", {}).get("outside", 2.5))
        section.different_first_page_header_footer = True
        enable_mirror_margins(section)

    # 2. Identifikacija odstavaka i uklanjanje nepravilnih
    paragraphs_to_remove = []
    section_type = "COVER_OUTER"
    skip_until_end = False

    for i, para in enumerate(doc.paragraphs):
        raw_text = para.text
        stripped_text = raw_text.strip()

        # Detekcija prelaza sekcija
        if section_type == "COVER_OUTER" and i > 15 and ("ALMA MATER EUROPAEA" in stripped_text or "Diplomski rad" in stripped_text):
            section_type = "COVER_INNER"
        elif section_type == "COVER_INNER" and ("ZAHVALNICA" in stripped_text or "SAŽETAK" in stripped_text or "ABSTRACT" in stripped_text):
            section_type = "FRONT_MATTER"
        elif stripped_text.startswith("1 UVOD") or stripped_text.startswith("1. UVOD"):
            section_type = "BODY"
        elif stripped_text in ["LITERATURA", "5 LITERATURA I IZVORI", "LITERATURA I IZVORI", "SEZNAM LITERATURE"]:
            section_type = "BIBLIOGRAPHY"
        elif stripped_text.startswith("PRILOG") or stripped_text.startswith("IZJAVA"):
            section_type = "APPENDIX"

        # Uklanjanje PRILOG C (Izvori grafikona) i PRILOG D (Izvori fotografija) po pravilu 18
        if "PRILOG C:" in stripped_text or "PRILOG D:" in stripped_text:
            skip_until_end = True
        elif skip_until_end and ("IZJAVA O AUTORSTVU" in stripped_text or "IZJAVA O LEKTORISANJU" in stripped_text):
            skip_until_end = False
            
        if skip_until_end:
            paragraphs_to_remove.append(para)
            continue

        # Uklanjanje sopstvenog vira po pravilu 15 (Vir: Anketni upitnik...)
        if (stripped_text.startswith("Vir:") or stripped_text.startswith("Izvor:")) and any(k in stripped_text.lower() for k in ["anketni", "lasten", "lastni", "avtor"]):
            paragraphs_to_remove.append(para)
            continue

        # Uklanjanje praznih redova u glavnom tekstu po pravilu 3
        if section_type == "BODY" and not stripped_text:
            paragraphs_to_remove.append(para)
            continue

    # Uklanjanje markiranih odstavaka iz dokumenta
    for p in paragraphs_to_remove:
        remove_paragraph(p)

    # 3. Drugi prolaz - stilizovanje preostalih odstavaka
    section_type = "COVER_OUTER"
    figures_captured = []
    uvod_paragraph = None

    for i, para in enumerate(doc.paragraphs):
        raw_text = para.text
        stripped_text = raw_text.strip()

        if section_type == "COVER_OUTER" and i > 15 and ("ALMA MATER EUROPAEA" in stripped_text or "Diplomski rad" in stripped_text):
            section_type = "COVER_INNER"
        elif section_type == "COVER_INNER" and ("ZAHVALNICA" in stripped_text or "SAŽETAK" in stripped_text or "ABSTRACT" in stripped_text):
            section_type = "FRONT_MATTER"
        elif stripped_text.startswith("1 UVOD") or stripped_text.startswith("1. UVOD"):
            section_type = "BODY"
            if uvod_paragraph is None:
                uvod_paragraph = para
        elif stripped_text in ["LITERATURA", "5 LITERATURA I IZVORI", "LITERATURA I IZVORI", "SEZNAM LITERATURE"]:
            section_type = "BIBLIOGRAPHY"
        elif stripped_text.startswith("PRILOG") or stripped_text.startswith("IZJAVA"):
            section_type = "APPENDIX"

        # Korekcija stičnih pomišljaja i preimenovanja
        new_text = fix_page_ranges_in_text(raw_text)

        # Uklanjanje crtice iz naslova ustanove (Pravilo 1)
        if "ALMA MATER EUROPAEA -" in new_text:
            new_text = new_text.replace("ALMA MATER EUROPAEA -", "ALMA MATER EUROPAEA")

        # Zamena "Grafikon 1", "Graf 1" u "Slika 1", "Slika 2" u tekstu i naslovima (Pravilo 12)
        new_text = re.sub(r'\b(Grafikon|Graf|Pie chart)\s*(\d+)', r'Slika \2', new_text)
        new_text = re.sub(r'\(Grafikon\s*(\d+)\)', r'(Slika \1)', new_text)
        new_text = re.sub(r'\(Graf\s*(\d+)\)', r'(Slika \1)', new_text)

        # Ispravka naslova priloga (Pravilo 5)
        if stripped_text.startswith("PRILOG A:"):
            new_text = "a) Anketni upitnik za nastavnike karakternih plesova"
        elif stripped_text.startswith("PRILOG B:"):
            new_text = "b) Anketni upitnik za učenike"

        if para.text != new_text:
            para.text = new_text

        p_fmt = para.paragraph_format

        # Uklanjanje kurziva u SVIM odstavcima (Pravilo 8)
        for r in para.runs:
            if r.italic and not ("http://" in r.text or "https://" in r.text or "doi.org" in r.text):
                r.italic = False

        # ==================== SPECIFIČNO FORMATIRANJE PO SEKCIJAMA ====================

        if section_type == "COVER_OUTER" or section_type == "COVER_INNER":
            p_fmt.line_spacing = 1.15
            p_fmt.space_before = Pt(0)
            p_fmt.space_after = Pt(0)
            for r in para.runs:
                r.font.name = font_name

        elif section_type == "FRONT_MATTER":
            is_front_heading = stripped_text in ["ZAHVALNICA", "SAŽETAK", "ABSTRACT", "KAZALO VSEBINE", "Kazalo vsebine", "KAZALO TABEL", "KAZALO SLIK", "Kazalo slik"]
            if is_front_heading:
                p_fmt.line_spacing = 1.15
                p_fmt.space_before = Pt(12)
                p_fmt.space_after = Pt(6)
                p_fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if stripped_text in ["KAZALO VSEBINE", "Kazalo vsebine"]:
                    para.text = "Kazalo vsebine"
                for r in para.runs:
                    r.font.name = font_name
                    r.font.size = Pt(14)
                    r.bold = True
            else:
                p_fmt.line_spacing = 1.15
                p_fmt.space_before = Pt(0)
                p_fmt.space_after = Pt(0) if "KAZALO" in stripped_text.upper() or "SLIKA" in stripped_text.upper() else Pt(12)
                p_fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for r in para.runs:
                    r.font.name = font_name
                    r.font.size = Pt(12)

        elif section_type == "BODY":
            is_heading = para.style.name.startswith("Heading") or para.style.name.startswith("Naslov") or re.match(r'^\d+(\.\d+)*\s+[A-ZČĆŽŠĐ]', stripped_text)
            is_caption = stripped_text.startswith("Slika ") or stripped_text.startswith("Tabela ")
            is_source = stripped_text.startswith("Vir:") or stripped_text.startswith("Izvor:")

            if is_heading:
                heading_match = re.match(r'^(\d+(\.\d+)*)', stripped_text)
                level = 1
                if heading_match:
                    dots = heading_match.group(1).count('.')
                    level = dots + 1

                p_fmt.line_spacing = 1.15
                p_fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT

                if level == 1:
                    # Nivo 1 (1 UVOD): 14pt BOLD UPPERCASE, prelom strane
                    p_fmt.page_break_before = True
                    p_fmt.space_before = Pt(12)
                    p_fmt.space_after = Pt(6)
                    para.text = new_text.upper()
                    for r in para.runs:
                        r.font.name = font_name
                        r.font.size = Pt(14)
                        r.bold = True
                        r.italic = False
                elif level == 2:
                    # Nivo 2 (1. nivo podpoglavja - npr. 2.1): 12pt BOLD po AMEU pravilniku (strana 6)
                    p_fmt.space_before = Pt(12)
                    p_fmt.space_after = Pt(0)
                    for r in para.runs:
                        r.font.name = font_name
                        r.font.size = Pt(12)
                        r.bold = True
                        r.italic = False
                else:
                    # Nivo 3 i niže (2. nivo podpoglavja - npr. 2.1.1, 3.2.1): 12pt NORMAL (NE-KREPKO) po pravilu 9 recenzenta!
                    p_fmt.space_before = Pt(12)
                    p_fmt.space_after = Pt(0)
                    for r in para.runs:
                        r.font.name = font_name
                        r.font.size = Pt(12)
                        r.bold = False
                        r.italic = False

            elif is_caption:
                # Naslov slike/tabele: SREDINSKA PORAVNAVA, Krepko 12pt, keep_with_next (Pravilo 10, 14)
                p_fmt.line_spacing = 1.15
                p_fmt.space_before = Pt(6)
                p_fmt.space_after = Pt(0)
                p_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_fmt.keep_with_next = True
                for r in para.runs:
                    r.font.name = font_name
                    r.font.size = Pt(12)
                    r.bold = True

                if stripped_text.startswith("Slika ") and stripped_text not in figures_captured:
                    figures_captured.append(stripped_text)

            elif is_source:
                # Vir: SREDINSKA PORAVNAVA, Navadno 12pt (Pravilo 11, 14)
                p_fmt.line_spacing = 1.15
                p_fmt.space_before = Pt(0)
                p_fmt.space_after = Pt(12)
                p_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_fmt.keep_with_next = False
                for r in para.runs:
                    r.font.name = font_name
                    r.font.size = Pt(12)
                    r.bold = False

            else:
                # Običan odstavak glavnog besedila: 1.15 razmik, 0pt pre, 12pt posle, obostrano (Pravilo 3)
                p_fmt.line_spacing = 1.15
                p_fmt.space_before = Pt(0)
                p_fmt.space_after = Pt(12)
                p_fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for r in para.runs:
                    r.font.name = font_name
                    r.font.size = Pt(12)

        elif section_type == "BIBLIOGRAPHY":
            if stripped_text in ["LITERATURA", "5 LITERATURA I IZVORI", "LITERATURA I IZVORI", "SEZNAM LITERATURE"]:
                p_fmt.line_spacing = 1.15
                p_fmt.space_before = Pt(12)
                p_fmt.space_after = Pt(6)
                p_fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_fmt.page_break_before = True
                para.text = "LITERATURA"
                for r in para.runs:
                    r.font.name = font_name
                    r.font.size = Pt(14)
                    r.bold = True
            else:
                # Literatura stavke: 1.15 razmik, 0pt pre, 0pt posle, obostrano (Pravilo 20)
                p_fmt.line_spacing = 1.15
                p_fmt.space_before = Pt(0)
                p_fmt.space_after = Pt(0)
                p_fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for r in para.runs:
                    r.font.name = font_name
                    r.font.size = Pt(12)

        elif section_type == "APPENDIX":
            if stripped_text in ["PRILOGE", "PRILOGA"]:
                p_fmt.line_spacing = 1.15
                p_fmt.space_before = Pt(12)
                p_fmt.space_after = Pt(6)
                p_fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_fmt.page_break_before = True
                para.text = "PRILOGE"
                for r in para.runs:
                    r.font.name = font_name
                    r.font.size = Pt(14)
                    r.bold = True
            elif stripped_text.startswith("IZJAVA"):
                p_fmt.line_spacing = 1.15
                p_fmt.space_before = Pt(12)
                p_fmt.space_after = Pt(6)
                p_fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_fmt.page_break_before = True
                for r in para.runs:
                    r.font.name = font_name
                    r.font.size = Pt(14)
                    r.bold = True
            else:
                p_fmt.line_spacing = 1.15
                p_fmt.space_before = Pt(0)
                p_fmt.space_after = Pt(0)
                p_fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for r in para.runs:
                    r.font.name = font_name
                    r.font.size = Pt(12)

    # 4. Formatiranje svih tabela u dokumentu
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                for p in cell.paragraphs:
                    p.paragraph_format.line_spacing = 1.15
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    if row_idx == 0:
                        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.font.name = font_name
                        r.font.size = Pt(12)

    # 5. Ubacivanje Kazala vsebine i Kazala slik pre UVOD-a
    has_toc = any("KAZALO VSEBINE" in p.text.upper() or "KAZALO" == p.text.strip().upper() for p in doc.paragraphs)
    if not has_toc and uvod_paragraph is not None:
        toc_heading = uvod_paragraph.insert_paragraph_before("Kazalo vsebine")
        toc_heading.paragraph_format.page_break_before = True
        toc_heading.paragraph_format.line_spacing = 1.15
        toc_heading.paragraph_format.space_before = Pt(12)
        toc_heading.paragraph_format.space_after = Pt(6)
        for r in toc_heading.runs:
            r.font.name = font_name
            r.font.size = Pt(14)
            r.bold = True

    has_lof = any("KAZALO SLIK" in p.text.upper() for p in doc.paragraphs)
    if not has_lof and uvod_paragraph is not None:
        lof_heading = uvod_paragraph.insert_paragraph_before("Kazalo slik")
        lof_heading.paragraph_format.page_break_before = True
        lof_heading.paragraph_format.line_spacing = 1.15
        lof_heading.paragraph_format.space_before = Pt(12)
        lof_heading.paragraph_format.space_after = Pt(6)
        for r in lof_heading.runs:
            r.font.name = font_name
            r.font.size = Pt(14)
            r.bold = True

        for fig_text in figures_captured:
            p_fig = uvod_paragraph.insert_paragraph_before(fig_text)
            p_fig.paragraph_format.line_spacing = 1.15
            p_fig.paragraph_format.space_before = Pt(0)
            p_fig.paragraph_format.space_after = Pt(0)
            p_fig.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for r in p_fig.runs:
                r.font.name = font_name
                r.font.size = Pt(12)

    doc.save(docx_output_path)
    print(f"Uspešno procesiran i sačuvan dokument: {docx_output_path}")

if __name__ == "__main__":
    input_file = "sample.docx"
    output_file = "sample-formatted.docx"
    format_academic_paper(input_file, output_file)