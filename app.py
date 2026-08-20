"""Streamlit UI za styleguard.

Dve celine: formatiranje dokumenta i upravljanje bibliotekom pravila.

Korak pregleda pravila pre primene nije ukras -- ekstrakcija iz pravilnika je
najnepouzdaniji deo lanca, pa se svako pravilo prikazuje uz poreklo, pouzdanost
i citat iz izvora, i ništa se ne primenjuje dok korisnik ne potvrdi.

Izgled prati „Classical" sistem iz redizajna: navigacija je traka na vrhu,
sporedni tekst stoji u uskoj koloni uz sadržaj, a akcenat je uvek potez i
podvlaka, nikad ispuna. Streamlit-ovi widgeti ostaju isti -- menja se samo
kako izgledaju, preko `_inject_custom_css()` i malih HTML fragmenata.
"""

from __future__ import annotations

import html
import io
import tempfile
from collections import Counter
from contextlib import contextmanager
from enum import Enum
from pathlib import Path

import docx
import streamlit as st
import streamlit.components.v1 as components
from docx.oxml.ns import qn

from styleguard import __version__, i18n, identity
from styleguard.i18n import t
from styleguard.extract.mate_client import MateClient, MateConfig
from styleguard.extract.pipeline import extract_rule_set, identify_institution
from styleguard.extract.source import NoTextLayerError, read_rules_document
from styleguard.formatting.engine import FormatOptions, format_document
from styleguard.library import RuleLibrary, suggest_display_name
from styleguard.search import search as search_rule_sets
from styleguard.rules import (
    Casing,
    Evidence,
    RuleSet,
    enum_type_for_path,
    field_type_for_path,
    iter_field_paths,
    load_rule_set,
    options_for_path,
    set_by_path,
)

PRESETS_DIR = Path(__file__).resolve().parent / "presets"
ASSETS_DIR = Path(__file__).resolve().parent / "assets"
REPO_URL = "https://github.com/antiv/styleguard"

NAV_PAGES = ("format", "library", "help")

# Visina skrolujućih lista, u pikselima. Dovoljno za šest-sedam redova: lista
# se vidi kao lista, a ne kao strana.
_LIST_HEIGHT = 300
_LIST_ROWS = 5
_TABLE_HEIGHT = 470
_TABLE_ROWS = 8

st.set_page_config(
    page_title="StyleGuard",
    page_icon=str(ASSETS_DIR / "icon-192.png"),
    layout="wide",
    initial_sidebar_state="collapsed",
)


# --------------------------------------------------------------------------
# Sitni prikazivači
# --------------------------------------------------------------------------


def esc(value) -> str:
    """Sve što ulazi u HTML fragment prolazi ovuda -- naslovi setova i citati
    iz pravilnika su tuđi tekst."""
    return html.escape(str(value), quote=True)


def _html(markup: str, container=None) -> None:
    (container or st).markdown(markup, unsafe_allow_html=True)


def _icon(name: str, size: int = 15) -> str:
    """Lucide, potez 1.5, nikad ispunjen -- emodži je dizajn izbacio.

    Globus i lupa stoje u CSS-u kao maska, jer vise uz Streamlit-ove widgete
    do kojih HTML odavde ne dopire.
    """
    paths = {
        "file-text": '<path d="M15 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V7Z"></path><path d="M14 2v5h6"></path>',
    }
    return (
        f'<svg class="sg-icon" width="{size}" height="{size}" viewBox="0 0 24 24" '
        f'fill="none" stroke="currentColor" stroke-width="1.5" stroke-linecap="round" '
        f'stroke-linejoin="round">{paths[name]}</svg>'
    )


def _mark_svg(size: int = 22) -> str:
    """Znak aplikacije u bojama sistema: tamna korica, list boje strane, zlatna hrbat-traka."""
    return (
        f'<svg viewBox="0 0 512 512" width="{size}" height="{size}" class="sg-mark">'
        '<rect width="512" height="512" rx="112" fill="#201f1d"></rect>'
        '<rect x="112" y="80" width="288" height="368" rx="24" fill="#f3f2f2"></rect>'
        '<rect x="176" y="168" width="108" height="16" rx="8" fill="#bab6b6"></rect>'
        '<rect x="190" y="210" width="152" height="16" rx="8" fill="#bab6b6"></rect>'
        '<rect x="172" y="252" width="124" height="16" rx="8" fill="#bab6b6"></rect>'
        '<rect x="160" y="300" width="200" height="16" rx="8" fill="#201f1d"></rect>'
        '<rect x="160" y="338" width="200" height="16" rx="8" fill="#201f1d"></rect>'
        '<rect x="160" y="376" width="200" height="16" rx="8" fill="#201f1d"></rect>'
        '<rect x="160" y="414" width="128" height="16" rx="8" fill="#201f1d"></rect>'
        '<rect x="136" y="160" width="10" height="270" rx="5" fill="#b68235"></rect>'
        "</svg>"
    )


def _kicker(text: str, muted: bool = False) -> str:
    css = "sg-kicker sg-kicker-muted" if muted else "sg-kicker"
    return f'<div class="{css}">{esc(text)}</div>'


def _step_head(number: str, title: str) -> str:
    return (
        f'<div class="sg-step"><span class="sg-step-no num">{esc(number)}</span>'
        f'<h4 class="sg-h4">{esc(title)}</h4></div>'
    )


def _human_size(size: int) -> str:
    if size < 1024:
        return f"{size} B"
    if size < 1024 * 1024:
        return f"{size / 1024:.1f} KB"
    return f"{size / (1024 * 1024):.1f} MB"


# --------------------------------------------------------------------------
# Zajedničko
# --------------------------------------------------------------------------


@st.cache_resource
def get_library() -> RuleLibrary:
    return RuleLibrary()


def load_presets() -> list[RuleSet]:
    out: list[RuleSet] = []
    for path in sorted(PRESETS_DIR.glob("*.json")):
        try:
            out.append(load_rule_set(path))
        except Exception:
            continue
    return out


@st.cache_data(ttl=600, show_spinner=False)
def _docx_facts(data: bytes) -> tuple[int, int] | None:
    """Broj pasusa i tabela za priloženi rad.

    Keširano po sadržaju: Streamlit preiscrtava stranu na svaki klik, a
    otvaranje dokumenta pri svakom prolazu bi se osećalo.
    """
    try:
        document = docx.Document(io.BytesIO(data))
    except Exception:
        return None
    return len(document.paragraphs), len(document.tables)


def _document_facts(document) -> dict[str, int]:
    """Brojevi kojima se izveštaj potkrepljuje -- čitaju se iz rezultata, ne iz glave."""
    body = document.element.body
    return {
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "images": len(body.findall(".//" + qn("w:drawing")))
        + len(body.findall(".//" + qn("w:pict"))),
    }


def mate_client() -> MateClient | None:
    config = MateConfig.from_env()
    return MateClient(config) if config.is_configured else None


@contextmanager
def temp_upload(uploaded):
    """Uploadovani fajl na disku samo dok traje obrada.

    Bitno na serveru: bez brisanja bi se tuđi pravilnici gomilali u /tmp i
    ostajali tamo neograničeno. Sam rad nikad ne dodiruje disk -- ide kroz
    `io.BytesIO`.
    """
    suffix = Path(uploaded.name).suffix
    handle = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    try:
        handle.write(uploaded.getbuffer())
        handle.close()
        yield Path(handle.name)
    finally:
        Path(handle.name).unlink(missing_ok=True)


@st.cache_data(ttl=120, show_spinner=False)
def _mate_ping(base_url: str, agent: str, token_fingerprint: str) -> tuple[bool, str]:
    """Keširan ping.

    Bez keša bi se Mate zvao pri svakom preiscrtavanju strane, a Streamlit
    preiscrtava na svaki klik. Argumenti ulaze u ključ keša da bi izmena
    konfiguracije oborila keš; token se ne prosleđuje u čitavom obliku.
    """
    return MateClient(MateConfig.from_env()).ping()


def mate_status_banner() -> None:
    """Status agenta kao obrisana pilula u zaglavlju.

    Ista provera kao ranije (isti keš), samo bez emodžija: stanje nosi tačka,
    ne boja slova.
    """
    config = MateConfig.from_env()
    if not config.is_configured:
        _html(
            f'<div class="sg-pill" title="{esc(t("mate.heuristic_only"))}">'
            f'<span class="sg-dot sg-dot-muted"></span>{esc(t("mate.heuristic_short"))}</div>'
        )
        return
    ok, message = _mate_ping(config.base_url, config.agent, (config.token or "")[-6:])
    state = "sg-dot-online" if ok else "sg-dot-warn"
    label = t("mate.agent_online") if ok else t("mate.agent_offline")
    _html(
        f'<div class="sg-pill" title="{esc(message)}">'
        f'<span class="sg-dot {state}"></span>{esc(config.agent)} · {esc(label)}</div>'
    )


# --------------------------------------------------------------------------
# Editor pravila
# --------------------------------------------------------------------------


def _source_label(source: str) -> str:
    return t(f"source.{source}")


_CONFIDENCE_STRENGTH = {"high": "", "medium": " sg-dot-medium", "low": " sg-dot-low"}


def _evidence_html(evidence: Evidence | None, missing: bool = False) -> str:
    """Poreklo vrednosti: zlatna tačka i citat, bez semafora u bojama.

    Ručno unesena vrednost nema citat -- tu stoji samo oznaka izvora, jer je
    korisnik izvor. „Ne pominje se u pravilniku" stoji samo tamo gde je polje
    zaista ostalo nerazrešeno; set bez zabeleženih citata nije isto što i
    pravilnik koji ćuti.
    """
    if evidence is None:
        text = t("rules.no_evidence") if missing else t("editor.none")
        return f'<div class="sg-ev sg-ev-empty"><span>{esc(text)}</span></div>'
    if evidence.source == "manual" or not evidence.quote:
        return f'<div class="sg-ev sg-ev-empty"><span>{esc(_source_label(evidence.source))}</span></div>'

    dot = _CONFIDENCE_STRENGTH.get(evidence.confidence, "")
    page = (
        f' <span class="sg-ev-page num">{esc(t("evidence.page", page=evidence.page))}</span>'
        if evidence.page
        else ""
    )
    return (
        f'<div class="sg-ev"><span class="sg-ev-dot{dot}"></span>'
        f"<span><em>{esc(evidence.quote)}</em>{page}</span></div>"
    )


# Grupa preko ovoliko polja se cepa na podgrupe po drugom segmentu putanje.
# `captions` ima dvadeset i četiri polja, a `typography` dva -- jedna kartica
# preko cele strane pored kartice od dva reda je ono što mrežu razbija.
_CARD_SPLIT_AT = 12


def _rule_cards(rule_set: RuleSet) -> list[tuple[str, list[tuple[str, object]]]]:
    """Polja raspoređena po karticama.

    Grupisanje je i dalje po prvom segmentu putanje; menja se samo raspored:
    prevelika grupa se deli na podgrupe, a kartice idu od najveće ka najmanjoj
    da bi susedi u redu bili slične visine. Redosled je određen brojem polja,
    dakle isti pri svakom prikazu istog seta.
    """
    groups: dict[str, list[tuple[str, object]]] = {}
    for path, value in iter_field_paths(rule_set.rules):
        groups.setdefault(path.split(".", 1)[0], []).append((path, value))

    cards: list[tuple[str, list[tuple[str, object]]]] = []
    for group, fields in groups.items():
        if len(fields) <= _CARD_SPLIT_AT:
            cards.append((group, fields))
            continue
        parts: dict[str, list[tuple[str, object]]] = {}
        for path, value in fields:
            rest = path.split(".", 1)[1] if "." in path else ""
            head = rest.split(".", 1)[0] if "." in rest else group
            parts.setdefault(f"{group}.{head}" if head != group else group, []).append(
                (path, value)
            )
        cards.extend(parts.items())

    order = {group: index for index, (group, _) in enumerate(groups.items())}
    return sorted(
        cards, key=lambda card: (-len(card[1]), order.get(card[0].split(".", 1)[0], 0))
    )


def rules_editor(rule_set: RuleSet, key_prefix: str) -> RuleSet:
    """Editor svih skalarnih polja, kartica po sekciji šeme.

    Grupisanje je isto kao pre -- prvi segment putanje. Promenio se samo sud:
    kartice u dve kolone umesto harmonike preko cele širine, jer se tako vidi
    koliko je pravilnik zaista rekao, a koliko je ostalo prazno.
    """
    unresolved = set(rule_set.unresolved)

    # Kartice u istom redu izjednačavaju visinu (vidi `sgcard` u CSS-u), pa se
    # susedi ne razilaze kad jedna grupa ima devet polja a druga dva.
    items = _rule_cards(rule_set)
    for index in range(0, len(items), 2):
        columns = st.columns(2, gap="large")
        for column, (group, fields) in zip(columns, items[index : index + 2]):
            with column:
                filled = sum(1 for _, value in fields if value is not None)
                with st.container(key=f"sgcard-{key_prefix}-{group}"):
                    _html(
                        '<div class="sg-card-head">'
                        f'<h4 class="sg-h4">{esc(group)}</h4>'
                        f'<span class="sg-card-count num">'
                        f"{esc(t('rules.group_count', filled=filled, total=len(fields)))}</span></div>"
                    )
                    for path, value in fields:
                        _render_rule_row(rule_set, path, value, unresolved, key_prefix, group)

    _render_headings_editor(rule_set, key_prefix)
    _render_keywords_editor(rule_set, key_prefix)
    return rule_set


def _render_rule_row(
    rule_set: RuleSet, path: str, value, unresolved: set[str], key_prefix: str, group: str
) -> None:
    """Jedan red kartice: naziv polja, vrednost, i rečenica iz koje je pročitana.

    Oznaka nosi samo ono što naslov kartice ne kaže -- u kartici `captions.figure`
    piše `position`, ne `figure.position`.
    """
    evidence = rule_set.evidence_for(path)
    label = path[len(group) + 1 :] if path.startswith(group + ".") else path
    is_missing = path in unresolved
    missing = " sg-field-missing" if is_missing else ""

    with st.container(key=f"sgrow-{key_prefix}-{path}"):
        columns = st.columns([0.62, 0.42, 1.0], vertical_alignment="center")
        _html(f'<div class="sg-field{missing}">{esc(label)}</div>', columns[0])

        options, enum_type = options_for_path(rule_set.rules, path)
        field_type = field_type_for_path(rule_set.rules, path)
        new_value = _render_input(
            path, value, f"{key_prefix}:{path}", columns[1],
            options=options, enum_type=enum_type, field_type=field_type
        )
        if new_value != value:
            set_by_path(rule_set.rules, path, new_value)
            rule_set.evidence = [e for e in rule_set.evidence if e.field_path != path]
            if new_value is not None:
                rule_set.evidence.append(
                    Evidence(field_path=path, confidence="high", source="manual")
                )
            rule_set.unresolved = [u for u in rule_set.unresolved if u != path]

        _html(_evidence_html(evidence, is_missing), columns[2])


def _render_input(path: str, value, key: str, column, options=None, enum_type=None, field_type=None):
    """Widget koji ume da vrati `None` -- prazno polje znači 'ne diraj'."""
    if options is not None:
        select_options = [t("editor.leave_alone")] + list(options)
        curr_str = value.value if isinstance(value, Enum) else (str(value) if value is not None else t("editor.leave_alone"))
        if curr_str not in select_options:
            curr_str = t("editor.leave_alone")
        chosen = column.selectbox(
            path, select_options, index=select_options.index(curr_str), key=key, label_visibility="collapsed"
        )
        if chosen == t("editor.leave_alone"):
            return None
        return enum_type(chosen) if enum_type is not None else chosen

    is_bool = (
        (field_type is not None and isinstance(field_type, type) and issubclass(field_type, bool))
        or isinstance(value, bool)
        or path.endswith(
            ("bold", "italic", "mirror_margins", "different_first_page", "allow_italic",
             "allow_empty_paragraphs", "insert_field", "header_row_bold", "header_row_repeat",
             "page_break_before", "keep_with_next")
        )
    )
    if is_bool:
        options_list = [t("editor.leave_alone"), t("editor.yes"), t("editor.no")]
        current = (
            t("editor.leave_alone")
            if value is None
            else (t("editor.yes") if value else t("editor.no"))
        )
        chosen = column.selectbox(
            path, options_list, index=options_list.index(current), key=key, label_visibility="collapsed"
        )
        return None if chosen == t("editor.leave_alone") else chosen == t("editor.yes")

    is_number = (
        (field_type is not None and isinstance(field_type, type) and issubclass(field_type, (int, float)))
        or isinstance(value, (int, float))
        or path.endswith(("_pt", "_cm", "spacing", "levels", "top", "bottom", "inside", "outside"))
    )
    if is_number:
        text = "" if value is None else str(value)
        entered = column.text_input(path, text, key=key, label_visibility="collapsed",
                                    placeholder=t("editor.placeholder"))
        if not entered.strip():
            return None
        try:
            return float(entered.replace(",", "."))
        except ValueError:
            column.error(t("editor.number_expected"))
            return value

    is_list = (
        (field_type is not None and isinstance(field_type, type) and issubclass(field_type, list))
        or isinstance(value, list)
    )
    if is_list:
        entered = column.text_input(
            path, ", ".join(str(v) for v in value) if isinstance(value, list) else (str(value) if value else ""), key=key, label_visibility="collapsed"
        )
        return [part.strip() for part in entered.split(",") if part.strip()]

    text = "" if value is None else str(value)
    entered = column.text_input(path, text, key=key, label_visibility="collapsed",
                                placeholder=t("editor.placeholder"))
    return entered.strip() or None


def _render_headings_editor(rule_set: RuleSet, key_prefix: str) -> None:
    with st.container(key=f"sgcard-{key_prefix}-headings"):
        _html(
            '<div class="sg-card-head">'
            f'<h4 class="sg-h4">{esc(t("editor.headings_title"))}</h4>'
            f'<span class="sg-card-count num">'
            f"{esc(t('rules.headings_count', count=len(rule_set.rules.headings)))}</span></div>"
        )
        if not rule_set.rules.headings:
            st.info(t("editor.headings_empty"))
            return

        # Zaglavlje kolona stoji jednom; ponovljene oznake uz svaki nivo su
        # četiri puta ista reč i pojedu red.
        with st.container(key=f"sgrow-{key_prefix}-headhead"):
            head = st.columns(6, vertical_alignment="center")
        for column, label in zip(
            head[1:],
            (t("editor.pt"), t("editor.bold"), t("editor.casing"),
             t("editor.page_break"), t("editor.keep_with_next")),
        ):
            _html(f'<div class="sg-colhead">{esc(label)}</div>', column)

        for heading in rule_set.rules.headings:
            with st.container(key=f"sgrow-{key_prefix}-h{heading.level}"):
                columns = st.columns(6, vertical_alignment="center")
                _html(
                    f'<div class="sg-field">{esc(t("editor.heading_level", level=heading.level))}</div>',
                    columns[0],
                )
                heading.size_pt = _float_input(columns[1], t("editor.pt"), heading.size_pt, f"{key_prefix}:h{heading.level}:size")
                heading.bold = _bool_input(columns[2], t("editor.bold"), heading.bold, f"{key_prefix}:h{heading.level}:bold")
                heading.casing = _casing_input(columns[3], heading.casing, f"{key_prefix}:h{heading.level}:casing")
                heading.page_break_before = _bool_input(
                    columns[4], t("editor.page_break"), heading.page_break_before, f"{key_prefix}:h{heading.level}:pb"
                )
                heading.keep_with_next = _bool_input(
                    columns[5], t("editor.keep_with_next"), heading.keep_with_next, f"{key_prefix}:h{heading.level}:kwn"
                )


def _render_keywords_editor(rule_set: RuleSet, key_prefix: str) -> None:
    keywords = rule_set.rules.structure_profile.section_keywords
    with st.container(key=f"sgcard-{key_prefix}-keywords"):
        _html(f'<div class="sg-card-head"><h4 class="sg-h4">{esc(t("editor.keywords"))}</h4></div>')
        st.caption(t("editor.keywords_help"))
        columns = st.columns(4)
        for column, field in zip(columns, ("front_matter", "body_start", "bibliography", "appendix")):
            current = getattr(keywords, field)
            entered = column.text_input(
                field, ", ".join(current), key=f"{key_prefix}:kw:{field}"
            )
            setattr(keywords, field, [p.strip() for p in entered.split(",") if p.strip()])


def _float_input(column, label, value, key):
    entered = column.text_input(
        label, "" if value is None else str(value), key=key, label_visibility="collapsed"
    )
    if not entered.strip():
        return None
    try:
        return float(entered.replace(",", "."))
    except ValueError:
        return value


def _bool_input(column, label, value, key):
    none_label = t("editor.none")
    options = [none_label, t("editor.yes"), t("editor.no")]
    current = none_label if value is None else (t("editor.yes") if value else t("editor.no"))
    chosen = column.selectbox(
        label, options, index=options.index(current), key=key, label_visibility="collapsed"
    )
    return None if chosen == none_label else chosen == t("editor.yes")


def _casing_input(column, value, key):
    none_label = t("editor.none")
    options = [none_label] + [member.value for member in Casing]
    current = value.value if value is not None else none_label
    chosen = column.selectbox(
        t("editor.casing"), options, index=options.index(current), key=key,
        label_visibility="collapsed",
    )
    return None if chosen == none_label else Casing(chosen)


# --------------------------------------------------------------------------
# Stranica: formatiranje
# --------------------------------------------------------------------------


def _set_active_rule_set(rule_set: RuleSet | None) -> None:
    """Postavlja aktivni set pravila i čisti stanje prethodnih widgeta u editoru."""
    for k in list(st.session_state.keys()):
        if k.startswith("fmt:"):
            del st.session_state[k]
    st.session_state["rule_set"] = rule_set
    st.session_state["rule_set_ver"] = st.session_state.get("rule_set_ver", 0) + 1
    # Editor se otvara podaleko ispod pregiba; bez ovoga klik na „Use" izgleda
    # kao da se ništa nije desilo.
    st.session_state["_scroll_to_rules"] = rule_set is not None


def _scroll_into_view(selector: str) -> None:
    """Spusti stranu na traženi deo.

    Ide kroz `components.html`, ne kroz `st.html`: ovaj drugi izbaci `<script>`
    iz dokumenta, pa se ništa ne izvrši. Komponenta je poseban okvir, otuda
    `window.parent.document`; sam okvir je nulte visine i sakriven CSS-om.

    Brojač menja sadržaj skripte pri svakom pozivu -- inače Streamlit prepozna
    isti element i ne pokrene ga ponovo, pa se druga akcija zaredom ne pomeri.
    Meta se traži u nekoliko pokušaja, jer skripta stiže u isti prolaz kao i
    sadržaj koji traži.
    """
    nonce = st.session_state.get("_scroll_nonce", 0) + 1
    st.session_state["_scroll_nonce"] = nonce
    components.html(
        f"<script>/*{nonce}*/(function(){{var d=window.parent.document;var n=0;"
        f"var f=function(){{var t=d.querySelector('{selector}');"
        "if(t){t.scrollIntoView({behavior:'smooth',block:'start'});}"
        "else if(++n<25){setTimeout(f,80);}};f();})();</script>",
        height=0,
    )


def _scroll_to_rules_once() -> None:
    """Jednom, odmah posle učitavanja, spusti stranu na pregled pravila."""
    if st.session_state.pop("_scroll_to_rules", False):
        _scroll_into_view(".sg-review-head")


def _thesis_facts(document_file) -> None:
    """Veličina, broj pasusa i tabela ispod priloženog rada.

    Ime fajla već nosi Streamlit-ov red sa dugmadima za uklanjanje i zamenu,
    pa se ovde ne ponavlja -- ostaje samo ono što taj red ne zna.
    """
    if document_file is None:
        return
    size = _human_size(len(document_file.getbuffer()))
    facts = _docx_facts(bytes(document_file.getbuffer()))
    meta = (
        t("format.file_facts", size=size, paragraphs=facts[0], tables=facts[1])
        if facts
        else size
    )
    _html(
        f'<div class="sg-file-meta num">{_icon("file-text", 13)} {esc(meta)}</div>'
    )


def _rules_step() -> str:
    """Korak 02: odakle dolaze pravila. Biblioteka je prva jer je najčešći put."""
    _html(_step_head("02", t("format.step.rules")))
    sources = {
        "library": t("format.source.library"),
        "upload": t("format.source.upload"),
        "json": t("format.source.json"),
    }
    with st.container(key="sg-seg"):
        source = st.pills(
            t("format.rules_source"),
            options=list(sources),
            default="upload",
            key="source_mode",
            format_func=lambda key: sources[key],
            label_visibility="collapsed",
        ) or "upload"

    if source == "upload":
        st.file_uploader(t("format.guide_file"), type=["pdf", "docx"], key="rules")
    elif source == "library":
        with st.container(key="sg-search"):
            st.text_input(
                t("library.search"),
                key="format:query",
                placeholder=t("library.search_placeholder"),
                label_visibility="collapsed",
            )
    else:
        st.file_uploader("rules.json", type=["json"], key="rules_json")
    return source


def _render_set_list(sets, bundled_ids, saved_count: int, key: str, on_use) -> None:
    """Lista setova umesto padajuće liste.

    Sa dvadesetak setova padajuća lista prestaje da bude izbor i postaje
    pretraživanje očima. Pretraga gleda naziv, id i instituciju, podnosi
    dijakritiku i padeže, i preslovljava ćirilicu -- „skopje" nalazi „Скопје".
    """
    query = st.session_state.get(f"{key}:query", "")
    found = search_rule_sets(sets, query)
    if not found:
        st.warning(t("library.no_match", query=query))
        return

    active = st.session_state.get("rule_set")
    active_id = active.meta.id if active is not None else None

    # Osamnaest redova preko cele strane potisne sve ostalo ispod pregiba, pa
    # duga lista dobija svoju visinu i skrol. Kratka ostaje kratka -- fiksna
    # visina bi ispod dva reda ostavila praznu kutiju.
    height = _LIST_HEIGHT if len(found) > _LIST_ROWS else "content"
    with st.container(key=f"sg-setlist-{key}", height=height):
        for rule_set in found:
            meta = rule_set.meta
            is_bundled = meta.id in bundled_ids
            where = " · ".join(
                part
                for part in (meta.institution.university, meta.institution.faculty)
                if part
            ) or "—"
            chosen = " sg-set-active" if meta.id == active_id else ""
            with st.container(key=f"sgset-{key}-{meta.id}"):
                columns = st.columns([1, 0.52, 0.34, 0.24], vertical_alignment="center")
                _html(
                    f'<div class="sg-set-name{chosen}">{esc(meta.display_name)}</div>'
                    f'<div class="sg-set-id num">{esc(meta.id)}</div>',
                    columns[0],
                )
                _html(f'<div class="sg-set-where">{esc(where)}</div>', columns[1])
                _html(
                    '<div><span class="sg-tag">'
                    f"{esc(t('library.bundled') if is_bundled else t('library.origin.saved'))}"
                    "</span></div>",
                    columns[2],
                )
                label = t("format.in_use") if meta.id == active_id else t("format.use")
                if columns[3].button(
                    label, key=f"use-{key}-{meta.id}", disabled=meta.id == active_id
                ):
                    on_use(rule_set)
    _html(
        '<div class="sg-set-foot num">'
        f"{esc(t('format.sets_footer', bundled=len(bundled_ids), saved=saved_count))}</div>"
    )


def page_format(user) -> None:
    library = get_library()

    saved = library.list()
    saved_ids = {rs.meta.id for rs in saved}
    bundled = [p for p in load_presets() if p.meta.id not in saved_ids]
    bundled_ids = {p.meta.id for p in bundled}
    # Jedna lista gotovih setova, ne dve. Odakle fajl dolazi -- `rules_library/`
    # ili `presets/` -- je detalj isporuke; korisnik pita samo da li pravila za
    # njegov fakultet već postoje.
    available = saved + bundled

    row = st.container(key="sg-hero-row")
    with row:
        hero, aside = st.columns([1, 0.26], gap="large")
    with aside:
        _render_marginalia()

    with hero:
        _html(f'<h1 class="sg-h1">{esc(t("format.header"))}</h1>')
        _html(f'<p class="sg-lead">{esc(t("format.intro"))}</p>')
        _html('<hr class="sg-hr">')

        with st.container(key="sg-steps"):
            left, right = st.columns(2, gap="large")
        with left:
            _html(_step_head("01", t("format.step.thesis")))
            picked = st.session_state.get("doc") is not None
            with st.container(key="sg-thesis-picked" if picked else "sg-thesis-empty"):
                document_file = st.file_uploader(
                    t("format.thesis"), type=["docx"], key="doc", label_visibility="collapsed"
                )
            _thesis_facts(document_file)
        with right:
            source = _rules_step()

        if source == "library":
            if not available:
                st.info(t("format.library_empty"))
            else:
                def use(chosen: RuleSet) -> None:
                    # Priložen set se kopira, ne učitava: izmene u editoru ne smeju
                    # da završe u fajlu koji je deo isporuke.
                    loaded = (
                        chosen.model_copy(deep=True)
                        if chosen.meta.id in bundled_ids
                        else library.load(chosen.meta.id)
                    )
                    _set_active_rule_set(loaded)
                    st.toast(t("format.loaded", name=loaded.meta.display_name))
                    st.rerun()

                _render_set_list(available, bundled_ids, len(saved), "format", use)

        _html('<hr class="sg-hr">')
        with st.container(key="sg-hero-actions", horizontal=True, vertical_alignment="center"):
            if source == "upload":
                rules_file = st.session_state.get("rules")
                is_extracting = st.session_state.get("is_extracting", False)
                if st.button(
                    t("format.extract"),
                    type="primary",
                    disabled=is_extracting or rules_file is None,
                    key="extract_rules_btn",
                ):
                    st.session_state["is_extracting"] = True
                    try:
                        extracted = _extract_flow(rules_file, library)
                        if extracted is not None:
                            _set_active_rule_set(extracted)
                    finally:
                        st.session_state["is_extracting"] = False
            elif source == "json":
                json_file = st.session_state.get("rules_json")
                if st.button(t("format.load"), type="primary", disabled=json_file is None):
                    loaded = RuleSet.model_validate_json(json_file.getvalue().decode("utf-8"))
                    _set_active_rule_set(loaded)
            _html(f'<div class="sg-hint">{esc(t("format.review_hint"))}</div>')

    rule_set: RuleSet | None = st.session_state.get("rule_set")
    if rule_set is None:
        return

    _render_rules_header(rule_set)
    _scroll_to_rules_once()
    rule_ver = st.session_state.get("rule_set_ver", 0)
    rule_set = rules_editor(rule_set, key_prefix=f"fmt:{rule_ver}")
    st.session_state["rule_set"] = rule_set

    _save_to_library_controls(rule_set, library, user)

    if document_file is None:
        st.info(t("format.need_docx"))
        return

    # Traka akcija se lepi za dno prozora: editor je dugačak, a „Formatiraj"
    # ne sme da bude nešto do čega se skroluje.
    with st.container(key="sg-apply", horizontal=True, vertical_alignment="center"):
        run = st.button(t("format.run"), type="primary")
        preview = st.button(t("format.preview_deletions"))
        cleanup = st.checkbox(t("format.clean_empty"), value=True)
        lenient = st.checkbox(t("format.lenient"), value=False)
        toc_options = {
            None: t("format.toc.by_rules"),
            True: t("format.toc.insert"),
            False: t("format.toc.skip"),
        }
        toc_choice = st.selectbox(
            t("format.toc"), list(toc_options), format_func=lambda key: toc_options[key]
        )

    options = FormatOptions(
        strict_structure=not lenient,
        clean_empty_paragraphs=cleanup,
        insert_toc=toc_choice,
    )
    if preview:
        _run_format(document_file, rule_set, FormatOptions(**{**options.__dict__, "dry_run": True}))
    if run:
        _run_format(document_file, rule_set, options, offer_download=True)


def _render_rules_header(rule_set: RuleSet) -> None:
    """Zaglavlje pregleda pravila: koliko je pročitano, koliko pravilnik ćuti.

    Broj nerazrešenih polja stoji kao cifra, pa žuto upozorenje više nije
    potrebno -- ista informacija, bez uzbune.
    """
    fields = list(iter_field_paths(rule_set.rules))
    filled = sum(1 for _, value in fields if value is not None)
    sources = Counter(e.source for e in rule_set.evidence)
    origin = _source_label(sources.most_common(1)[0][0]) if sources else _source_label("manual")

    _html(
        f'<h2 class="sg-sr-only">{esc(t("format.review_rules"))}</h2>'
        '<div class="sg-review-head">'
        "<div>"
        f'{_kicker(t("rules.kicker", source=origin))}'
        f'<h2 class="sg-h2">{esc(rule_set.meta.display_name)}</h2>'
        "</div>"
        '<div class="sg-figures">'
        f'<div class="sg-figure"><div class="sg-figure-value num">{filled}</div>'
        f'<div class="sg-figure-label">{esc(t("rules.count.set"))}</div></div>'
        f'<div class="sg-figure"><div class="sg-figure-value num sg-accent">{len(rule_set.unresolved)}</div>'
        f'<div class="sg-figure-label">{esc(t("rules.count.missing"))}</div></div>'
        "</div></div>"
        f'<p class="sg-lead sg-lead-narrow">{esc(t("rules.intro"))}</p>'
    )


def _extract_flow(rules_file, library: RuleLibrary) -> RuleSet | None:
    with st.status(t("extract.reading"), expanded=True) as status:
        try:
            with temp_upload(rules_file) as path:
                document = read_rules_document(path)
        except NoTextLayerError as exc:
            status.update(label=str(exc), state="error")
            st.error(str(exc))
            return None
        except ValueError as exc:
            status.update(label=str(exc), state="error")
            st.error(str(exc))
            return None

        client = mate_client()

        status.update(label=t("extract.identifying"), state="running")
        institution, origin = identify_institution(document, client)
        st.write(
            t(
                "extract.institution",
                origin=origin,
                name=suggest_display_name(institution) or t("extract.institution_unknown"),
            )
        )

        matches = library.find_matches(institution)
        if matches and (matches[0].is_strong or matches[0].is_suggestion):
            best = matches[0]
            st.session_state["pending_match"] = best.rule_set.meta.id
            st.info(
                t(
                    "extract.match_details",
                    name=best.rule_set.meta.display_name,
                    score=f"{best.score:.2f}",
                )
            )

        status.update(label=t("extract.extracting"), state="running")
        outcome = extract_rule_set(
            document, client=client, library=library, institution=institution,
            on_progress=lambda msg: st.write(msg),
        )
        status.update(label=t("extract.done", source=outcome.source), state="complete")

    for warning in outcome.warnings:
        st.warning(warning)
    if outcome.rejected:
        st.error(t("extract.rejected", fields=", ".join(outcome.rejected)))
    return outcome.rule_set


def _save_to_library_controls(rule_set: RuleSet, library: RuleLibrary, user) -> None:
    with st.expander(t("save.expander")):
        if not identity.can_create(user):
            st.caption(t("save.requires_login"))
            st.download_button(
                t("save.download_json"),
                rule_set.model_dump_json(indent=2),
                file_name=f"{rule_set.meta.id}.json",
                mime="application/json",
            )
            return

        # Postojeći tuđi set se ne prepisuje: snimanje pravi novi, u vlasništvu
        # onoga ko snima. Original ostaje netaknut.
        existing = library.load(rule_set.meta.id) if library.exists(rule_set.meta.id) else None
        overwriting = existing is not None and identity.can_edit(existing, user)
        if existing is not None and not overwriting:
            st.caption(t("save.foreign", name=existing.meta.display_name))

        columns = st.columns([4, 1], vertical_alignment="bottom")
        name = columns[0].text_input(t("save.name"), rule_set.meta.display_name, key="save_name")
        if columns[1].button(t("save.button"), width='stretch'):
            if name != rule_set.meta.display_name:
                rule_set.meta.display_name = name
            if not overwriting:
                rule_set.meta.id = library.unique_id(name)
            rule_set.meta.owner = user.email
            rule_set.meta.owner_name = user.name
            library.save(rule_set)
            st.success(t("save.saved", id=rule_set.meta.id))
        st.download_button(
            t("save.download_json"),
            rule_set.model_dump_json(indent=2),
            file_name=f"{rule_set.meta.id}.json",
            mime="application/json",
        )


def _changes_table(report) -> str:
    rows = "".join(
        "<tr>"
        f'<td>{esc(summary.role or "—")}</td>'
        f'<td class="sg-td-muted">{esc(summary.rule_path)}</td>'
        f'<td class="num sg-td-right">{summary.count}</td>'
        f'<td class="num">{esc(summary.describe_transitions())}</td>'
        "</tr>"
        for summary in report.by_rule()
    )
    return (
        '<table class="sg-table"><thead><tr>'
        f'<th>{esc(t("report.col.role"))}</th>'
        f'<th>{esc(t("report.col.property"))}</th>'
        f'<th class="sg-td-right">{esc(t("report.col.count"))}</th>'
        f'<th>{esc(t("report.col.changes"))}</th>'
        f"</tr></thead><tbody>{rows}</tbody></table>"
    )


def _deletions_box(report) -> str:
    rows = "".join(
        '<div class="sg-del-row">'
        f'<span class="num sg-del-index">¶ {esc(change.paragraph_index)}</span>'
        f'<span class="sg-del-section">{esc(change.section or "—")}</span>'
        f"<span>{esc(change.detail or change.rule_path)}</span></div>"
        for change in report.deletions
    )
    return (
        '<div class="sg-box"><div class="sg-card-head">'
        f'<h4 class="sg-h4">{esc(t("report.deletions_title"))}</h4>'
        f'<span class="sg-card-count num">'
        f"{esc(t('report.empty_paragraphs', count=len(report.deletions)))}</span></div>"
        f"{rows}</div>"
    )


def _run_format(document_file, rule_set: RuleSet, options: FormatOptions, offer_download: bool = False) -> None:
    try:
        result = format_document(io.BytesIO(document_file.getbuffer()), rule_set, options)
    except Exception as exc:
        st.error(t("format.failed", error=exc))
        return

    report = result.report
    _html('<div class="sg-report-anchor"></div>')
    with st.container(key="sg-report-row"):
        left, right = st.columns([1, 0.42], gap="large")

    with left:
        _html(f'<h2 class="sg-h2">{esc(t("report.header"))}</h2>')
        _html(f'<p class="sg-lead sg-lead-narrow">{esc(t("report.intro"))}</p>')
        figures = (
            (len(report.style_changes), t("report.style_changes")),
            (len(report.deletions), t("report.deletions")),
            (len(report.insertions), t("report.insertions")),
        )
        _html(
            '<div class="sg-figrow">'
            + "".join(
                f'<div class="sg-figure"><div class="sg-figure-big num">{value}</div>'
                f'<div class="sg-figure-label">{esc(label)}</div></div>'
                for value, label in figures
            )
            + "</div>"
        )
        if options.dry_run:
            st.info(t("report.dry_run"))
        for warning in report.warnings:
            st.warning(warning)
        if report.style_changes:
            _html(_changes_table(report))
        if report.deletions:
            _html(_deletions_box(report))
        if not report.changes:
            st.info(t("report.no_changes"))

    with right:
        payload = result.to_bytes() if offer_download else None
        name = f"Formatiran_{Path(document_file.name).stem}.docx"
        _html(
            _kicker(t("report.result.title"))
            + f'<div class="sg-result-name">{esc(name)}</div>'
            + (
                f'<div class="sg-result-meta num">'
                f"{esc(t('report.file_facts', size=_human_size(len(payload))))}</div>"
                if payload is not None
                else ""
            )
        )
        if payload is not None:
            with st.container(key="sg-download"):
                st.download_button(
                    t("format.download"),
                    payload,
                    file_name=name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary",
                    width='stretch',
                )
        _html('<hr class="sg-hr">')
        facts = _document_facts(result.document)
        lines = [t("report.invariants.paragraphs", count=facts["paragraphs"])]
        if facts["tables"]:
            lines.append(t("report.invariants.tables", count=facts["tables"]))
        if facts["images"]:
            lines.append(t("report.invariants.images", count=facts["images"]))
        lines.append(t("report.invariants.links"))
        _html(
            f'<p class="sg-aside-body">{esc(t("report.invariants.lead"))}</p>'
            '<div class="sg-checks">'
            + "".join(
                f'<div class="sg-check"><span class="sg-accent">✓</span><span>{esc(line)}</span></div>'
                for line in lines
            )
            + "</div>"
        )

    # Traka akcija stoji na dnu prozora, pa se izveštaj otvara ispod pregiba --
    # rezultat mora sam da dođe pred oči, a ne da se do njega skroluje.
    _scroll_into_view(".sg-report-anchor")


# --------------------------------------------------------------------------
# Stranica: biblioteka
# --------------------------------------------------------------------------


def page_library(user) -> None:
    library = get_library()

    # Priloženi preseti se prikazuju ovde zajedno sa snimljenim setovima. Bez
    # toga su nevidljivi: biblioteka je prvo mesto na kome ih korisnik traži, a
    # oni su dotad postojali samo kao izvor pravila na strani za formatiranje.
    saved = library.list()
    saved_ids = {rs.meta.id for rs in saved}
    bundled = [p for p in load_presets() if p.meta.id not in saved_ids]
    bundled_ids = {p.meta.id for p in bundled}
    sets = saved + bundled

    with st.container(key="sg-lib-head"):
        head, tools = st.columns([1, 0.44], gap="large")
    with head:
        _html(
            f'<h2 class="sg-h2">{esc(t("library.header"))}</h2>'
            f'<p class="sg-lead sg-lead-narrow">{esc(t("library.subtitle"))}</p>'
        )
    with tools:
        with st.container(key="sg-search"):
            st.text_input(
                t("library.search"),
                key="library:query",
                placeholder=t("library.search_count", count=len(sets)),
                label_visibility="collapsed",
            )
        if identity.can_create(user):
            # Uvoz stoji iza dugmeta, ne kao trajna zona za prevlačenje: koristi
            # se retko, a preko cele širine bi gurao tabelu naniže.
            with st.container(key="sg-import", width="content"):
                popover = st.popover(t("library.import_file"))
            with popover:
                uploaded = st.file_uploader(
                    t("library.import_file"), type=["json"], key="import_json",
                    label_visibility="collapsed",
                )
                if uploaded and st.button(t("library.import"), type="primary"):
                    with temp_upload(uploaded) as path:
                        imported = library.import_(path)
                    # Uvezen set pripada onome ko ga je uveo, bez obzira na to čiji je
                    # bio u fajlu -- inače bi uvoz mogao da podmetne tuđe vlasništvo.
                    imported.meta.owner = user.email
                    imported.meta.owner_name = user.name
                    library.save(imported)
                    st.success(t("library.imported", id=imported.meta.id))
                    st.rerun()

    if user is None:
        st.info(t("library.anonymous"))
    elif user.is_admin:
        st.caption(t("library.admin"))

    if not sets:
        st.info(t("library.empty"))
        return

    found = search_rule_sets(sets, st.session_state.get("library:query", ""))
    if not found:
        st.warning(t("library.no_match", query=st.session_state.get("library:query", "")))
        return

    _render_library_table(found, bundled_ids, user, library)
    _html(f'<div class="sg-set-foot">{esc(t("library.bundled_hint"))}</div>')

    editing = st.session_state.get("editing")
    if editing and library.exists(editing):
        rule_set = library.load(editing)
        # Dozvola se proverava i ovde, ne samo na dugmetu: `editing` preživi
        # u sesiji, pa odjava ili izbor tuđeg seta ne sme da ostavi otvoren
        # editor sa aktivnim snimanjem.
        if not identity.can_edit(rule_set, user):
            st.session_state.pop("editing", None)
            st.rerun()

        _html('<hr class="sg-hr">')
        _html(f'<h2 class="sg-h2">{esc(t("library.editing", name=rule_set.meta.display_name))}</h2>')

        institution = rule_set.meta.institution
        columns = st.columns(4)
        rule_set.meta.display_name = columns[0].text_input(
            t("library.name"), rule_set.meta.display_name
        )
        institution.university = columns[1].text_input(
            t("library.university"), institution.university or ""
        ) or None
        institution.faculty = columns[2].text_input(
            t("library.faculty"), institution.faculty or ""
        ) or None
        institution.document_type = columns[3].text_input(
            t("library.document_type"), institution.document_type or ""
        ) or None

        lib_ver = st.session_state.get("lib_edit_ver", 0)
        rule_set = rules_editor(rule_set, key_prefix=f"lib:{editing}:{lib_ver}")
        if st.button(t("library.save_changes"), type="primary"):
            library.save(rule_set)
            st.toast(t("library.saved"))
            st.rerun()


def _render_library_table(sets, bundled_ids, user, library: RuleLibrary) -> None:
    """Tabela setova sa akcijama u redu, umesto dataframe-a i četiri dugmeta ispod."""
    height = _TABLE_HEIGHT if len(sets) > _TABLE_ROWS else "content"
    with st.container(key="sg-libtable", height=height):
        # Zaglavlje stoji unutar skrol-oblasti i lepi se za njen vrh -- na uskom
        # ekranu tabela klizi vodoravno, pa imena kolona moraju da klize s njom.
        _html(
            '<div class="sg-table sg-table-head"><div class="sg-lib-row">'
            f'<span>{esc(t("library.col.name"))}</span>'
            f'<span>{esc(t("library.col.university"))}</span>'
            f'<span>{esc(t("library.col.faculty"))}</span>'
            f'<span>{esc(t("library.col.type"))}</span>'
            f'<span>{esc(t("library.col.owner"))}</span>'
            f'<span>{esc(t("library.col.updated"))}</span>'
            "<span></span></div></div>"
        )
        _render_library_rows(sets, bundled_ids, user, library)


def _render_library_rows(sets, bundled_ids, user, library: RuleLibrary) -> None:
    for rule_set in sets:
        meta = rule_set.meta
        # Preset je fajl isporučen uz aplikaciju, ne podatak instance: izmena bi
        # nestala pri sledećem redeployu, a brisanje bi se vratilo. Kopija je
        # jedini put koji vodi negde, pa je jedina ponuđena.
        is_bundled = meta.id in bundled_ids
        may_edit = identity.can_edit(rule_set, user) and not is_bundled
        owner = (
            t("library.bundled")
            if is_bundled
            else meta.owner_name or meta.owner or t("library.builtin")
        )
        with st.container(key=f"sglib-{meta.id}"):
            columns = st.columns(
                [1.15, 0.75, 0.75, 0.5, 0.6, 0.5, 1.15], vertical_alignment="center"
            )
            _html(f'<div class="sg-set-name">{esc(meta.display_name)}</div>', columns[0])
            _html(f'<div class="sg-td-muted">{esc(meta.institution.university or "—")}</div>', columns[1])
            _html(f'<div class="sg-td-muted">{esc(meta.institution.faculty or "—")}</div>', columns[2])
            _html(
                f'<div><span class="sg-tag">{esc(meta.institution.document_type or "—")}</span></div>',
                columns[3],
            )
            _html(f'<div class="sg-td-muted">{esc(owner)}</div>', columns[4])
            _html(
                f'<div class="sg-td-muted num">{esc(meta.updated_at.strftime("%Y-%m-%d"))}</div>',
                columns[5],
            )
            with columns[6]:
                with st.container(horizontal=True, key=f"sglibact-{meta.id}"):
                    if st.button(t("library.edit"), key=f"edit-{meta.id}", type="tertiary", disabled=not may_edit):
                        for k in list(st.session_state.keys()):
                            if k.startswith("lib:"):
                                del st.session_state[k]
                        st.session_state["editing"] = meta.id
                        st.session_state["lib_edit_ver"] = st.session_state.get("lib_edit_ver", 0) + 1
                        st.rerun()
                    # Kopiranje je namerno dozvoljeno svakome ko je prijavljen, i nad tuđim
                    # setom: to je izlaz koji zabranu izmene čini neblokirajućom.
                    if st.button(
                        t("library.copy"), key=f"copy-{meta.id}", type="tertiary",
                        disabled=not identity.can_create(user),
                    ):
                        copy = library.duplicate_of(rule_set)
                        copy.meta.owner = user.email
                        copy.meta.owner_name = user.name
                        library.save(copy)
                        st.toast(t("library.copied", id=copy.meta.id))
                        st.rerun()
                    if st.button(t("library.delete"), key=f"del-{meta.id}", type="tertiary", disabled=not may_edit):
                        library.delete(meta.id)
                        st.session_state.pop("editing", None)
                        st.toast(t("library.saved"))
                        st.rerun()
                    st.download_button(
                        t("library.export"),
                        rule_set.model_dump_json(indent=2),
                        file_name=f"{meta.id}.json",
                        mime="application/json",
                        key=f"exp-{meta.id}",
                        type="tertiary",
                    )


# --------------------------------------------------------------------------
# Stranica: pomoć
# --------------------------------------------------------------------------


def page_help() -> None:
    """Opis aplikacije, uputstvo i šta se dešava sa dokumentima."""
    _html(
        f'<h1 class="sg-h1">{esc(t("help.header"))}</h1>'
        f'<p class="sg-lead sg-lead-wide">{esc(t("help.intro"))}</p>'
    )

    _html('<hr class="sg-hr sg-hr-tight">')
    with st.container(key="sg-help-cols"):
        columns = st.columns(3, gap="medium")
    with columns[0]:
        with st.container(key="sg-helpcol-1"):
            _html(_kicker(t("help.kicker.how")) + f'<h4 class="sg-h4">{esc(t("help.how.title"))}</h4>')
            with st.container(key="sg-justify-1"):
                st.markdown(t("help.how.body"))
    with columns[1]:
        with st.container(key="sg-helpcol-2"):
            _html(_kicker(t("help.kicker.documents")) + f'<h4 class="sg-h4">{esc(t("help.privacy.title"))}</h4>')
            with st.container(key="sg-justify-2"):
                st.markdown(t("help.privacy.lead"))
                st.markdown(f"**{t('help.privacy.thesis.title')}**")
                st.markdown(t("help.privacy.thesis.body"))
                st.markdown(f"**{t('help.privacy.guide.title')}**")
                st.markdown(t("help.privacy.guide.body"))
                st.markdown(f"**{t('help.privacy.stored.title')}**")
                st.markdown(t("help.privacy.stored.body"))
                if not MateConfig.from_env().is_configured:
                    st.info(t("help.privacy.offline"))
    with columns[2]:
        with st.container(key="sg-helpcol-3"):
            _html(_kicker(t("help.kicker.never")) + f'<h4 class="sg-h4">{esc(t("help.invariants.title"))}</h4>')
            with st.container(key="sg-justify-3"):
                st.markdown(t("help.invariants.body"))
                st.markdown(f"**{t('help.library.title')}**")
                st.markdown(t("help.library.body"))

    _html('<hr class="sg-hr">')
    _html(f'<h4 class="sg-h4">{esc(t("help.faq.title"))}</h4>')
    with st.container(key="sg-faq"):
        faq = st.columns(2, gap="large")
    for number in range(1, 6):
        with faq[(number - 1) % 2]:
            _html(
                '<div class="sg-faq">'
                f'<div class="sg-faq-q">{esc(t(f"help.faq.q{number}"))}</div>'
                f'<p class="sg-faq-a">{esc(t(f"help.faq.a{number}"))}</p></div>'
            )

    _html('<hr class="sg-hr">')
    _html(
        f'<h4 class="sg-h4">{esc(t("help.opensource.title"))}</h4>'
        f'<p class="sg-lead sg-lead-narrow">{esc(t("help.opensource.body"))}</p>'
    )
    st.link_button(t("help.opensource.button"), REPO_URL)


# --------------------------------------------------------------------------
# Okvir aplikacije: traka na vrhu i sporedna kolona
# --------------------------------------------------------------------------


def _sync_language_client(chosen: str) -> None:
    """Pamti izabrani jezik u localStorage i učitava ga na refresh."""
    js = f"""
    <script>
    (function() {{
        try {{
            var saved = localStorage.getItem('styleguard_lang');
            var active = "{chosen}";
            var url = new URL(window.location.href);
            var currentParam = url.searchParams.get('lang');

            // Ako u localStorage postoji sačuvan jezik, a trenutno aktivni jezik na serveru nije taj:
            if (saved && ['sr', 'en', 'fr', 'de'].indexOf(saved) !== -1 && saved !== active && !currentParam) {{
                url.searchParams.set('lang', saved);
                window.location.replace(url.href);
                return;
            }}

            // Uvek sinhronizuj localStorage sa aktivnim izborom:
            if (active) {{
                localStorage.setItem('styleguard_lang', active);
            }}
        }} catch(e) {{
            console.warn('localStorage error', e);
        }}
    }})();
    </script>
    """
    st.html(js, unsafe_allow_javascript=True)


def _apply_language() -> None:
    """Set the language for this session before anything is rendered.

    Precedence:
    1. Active choice in `session_state` (user changed in dropdown)
    2. URL query param `lang` (from link or localStorage sync)
    3. Browser `Accept-Language` HTTP header
    """
    valid = i18n.available_languages()
    chosen = st.session_state.get("language")

    if chosen not in valid:
        lang_param = st.query_params.get("lang")
        if isinstance(lang_param, list):
            lang_param = lang_param[0] if lang_param else None

        if lang_param and lang_param in valid:
            chosen = lang_param
        else:
            try:
                header = st.context.headers.get("Accept-Language")
            except Exception:
                header = None
            chosen = i18n.negotiate(header)

        st.session_state["language"] = chosen

    i18n.set_language(chosen)
    _sync_language_client(chosen)


def _on_language_change(picked: str) -> None:
    if picked in i18n.available_languages():
        st.session_state["language"] = picked
        st.query_params["lang"] = picked
        i18n.set_language(picked)


def _language_selector() -> None:
    """Globus i kôd jezika; spisak se otvara kao meni, ne kao padajuća lista preko cele širine."""
    current = st.session_state.get("language", i18n.DEFAULT_LANGUAGE)
    with st.container(key="sg-lang", width="content"):
        with st.popover(current.upper()):
            for code in i18n.available_languages():
                st.button(
                    i18n.language_name(code),
                    key=f"lang-{code}",
                    type="tertiary",
                    width="stretch",
                    on_click=_on_language_change,
                    args=(code,),
                )


def _go(page: str) -> None:
    st.session_state["nav_page"] = page


def _app_header(user, page: str) -> None:
    """Traka na vrhu: znak, navigacija, status agenta, jezik, nalog."""
    with st.container(key="sg-header", horizontal=True, vertical_alignment="center"):
        with st.container(key="sg-brand", horizontal=True, vertical_alignment="center", width="content"):
            _html(f'<div class="sg-brand">{_mark_svg(22)}<span>StyleGuard</span></div>')
        with st.container(key="sg-nav", horizontal=True, vertical_alignment="center", width="content"):
            for code in NAV_PAGES:
                st.button(
                    t(f"nav.{code}"),
                    key=f"nav-{code}",
                    type="tertiary",
                    on_click=_go,
                    args=(code,),
                )
        with st.container(key="sg-tools", horizontal=True, vertical_alignment="center", width="content"):
            mate_status_banner()
            _language_selector()
            identity.render_account(st, user)


def _render_marginalia() -> None:
    """Uska kolona uz sadržaj: obećanje aplikacije i stanje ove instance.

    Ranije je stajalo u sidebar-u kao niz caption-a; tu je smetalo navigaciji,
    a ovde je ono što jeste -- napomena sa strane.
    """
    instance: list[str] = []
    if not identity.oidc_configured(st):
        instance.append(t("auth.local_mode"))
    if not MateConfig.from_env().is_configured:
        instance.append(t("help.privacy.offline"))

    parts = [
        '<aside class="sg-aside">',
        "<div>",
        _kicker(t("aside.promise.title")),
        f'<p class="sg-aside-body">{esc(t("app.promise"))}</p>',
        "</div>",
    ]
    if instance:
        parts += [
            '<hr class="sg-hr sg-hr-tight">',
            "<div>",
            _kicker(t("aside.instance.title"), muted=True),
            f'<p class="sg-aside-body sg-aside-dim">{esc(" ".join(instance))}</p>',
            "</div>",
        ]
    parts += [
        f'<div class="sg-aside-foot num">v{esc(__version__)} · MIT · '
        f'<a href="{REPO_URL}">{esc(t("aside.github"))}</a></div>',
        "</aside>",
    ]
    _html("".join(parts))


# --------------------------------------------------------------------------
# Izgled
# --------------------------------------------------------------------------

_CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Cormorant+Garamond:wght@400;600&family=Lora:wght@400;600&display=swap');

:root {
  --sg-bg: #f3f2f2;
  --sg-surface: #eae9e9;
  --sg-text: #201f1d;
  --sg-accent: #b68235;
  --sg-accent-700: #7d5411;
  --sg-divider: rgba(32,31,29,0.16);
  --sg-n100: #f8f4f4;
  --sg-n500: #9b9797;
  --sg-n800: #444141;
  --sg-heading: "Cormorant Garamond", Georgia, serif;
  --sg-body: "Lora", Georgia, serif;
  --sg-radius: 4px;
  /* Bočni odmak strane. Traka na vrhu i traka akcija idu od ivice do ivice
     preko negativne margine, pa moraju da čitaju istu vrednost. */
  --sg-pad: 40px;
  --sg-shadow-sm: 0 1px 2px rgba(45,43,43,0.14);
  --sg-shadow-md: 0 3px 10px rgba(45,43,43,0.16);
  --sg-search: url('data:image/svg+xml;utf8,<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 24 24" fill="none" stroke="black" stroke-width="1.5"><circle cx="11" cy="11" r="7"/><path d="m20 20-4.3-4.3"/></svg>');
  --sg-globe: url('data:image/svg+xml;utf8,<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 24 24" fill="none" stroke="black" stroke-width="1.5"><circle cx="12" cy="12" r="10"/><path d="M2 12h20"/><path d="M12 2a15 15 0 0 1 0 20 15 15 0 0 1 0-20"/></svg>');
}

/* — okvir strane — */
.stApp { background: var(--sg-bg); color: var(--sg-text); font-family: var(--sg-body); }
header[data-testid="stHeader"] { display: none !important; }
.stAppViewContainer > .main .block-container,
.stMainBlockContainer {
  padding: 0 var(--sg-pad) 120px !important;
  max-width: 1420px !important;
}
.stApp, .stMarkdown, .stMarkdown p, .stMarkdown li,
label, input, textarea, button, select { font-family: var(--sg-body); }
h1, h2, h3, h4, h5, h6, .sg-h1, .sg-h2, .sg-h4 {
  font-family: var(--sg-heading); letter-spacing: -0.015em; line-height: 1.12;
}
::selection { background: rgba(182,130,53,0.3); }
*:focus-visible {
  outline: 2px solid var(--sg-accent) !important;
  outline-offset: 2px !important;
  box-shadow: none !important;
}
.num { font-variant-numeric: tabular-nums; font-feature-settings: "tnum"; }
.sg-accent { color: var(--sg-accent-700); }
.sg-icon { display: inline-block; vertical-align: middle; }

/* — tipografija — */
.sg-h1 { font-size: 44px; font-weight: 400; margin: 0 0 10px; }
.sg-h2 { font-size: 33px; font-weight: 400; margin: 0 0 6px; }
.sg-h4 { font-size: 19px; font-weight: 600; margin: 0; }
.sg-lead {
  max-width: 56ch; font-size: 15px; line-height: 1.65; margin: 0 0 22px;
  color: color-mix(in srgb, var(--sg-text) 78%, transparent);
}
.sg-lead-narrow { max-width: 74ch; font-size: 14px; line-height: 1.6; }
.sg-lead-wide { max-width: 64ch; }
.sg-hint { font-size: 12px; color: color-mix(in srgb, var(--sg-text) 55%, transparent); }
.sg-hr { height: 1px; border: 0; margin: 22px 0; background: var(--sg-divider); }
.sg-hr-tight { margin: 14px 0; }
.sg-kicker {
  font-family: var(--sg-heading); font-size: 11px; letter-spacing: 0.12em;
  text-transform: uppercase; color: var(--sg-accent); margin-bottom: 7px;
}
.sg-kicker-muted { color: color-mix(in srgb, var(--sg-text) 55%, transparent); }

/* — traka na vrhu — */
.st-key-sg-header {
  border-bottom: 1px solid var(--sg-divider);
  padding: 11px var(--sg-pad) !important;
  /* Negativna margina pomera traku ulevo, ali joj ne menja širinu: bez
     izričite širine i ukinutog `max-width` desni prepust ostaje neispunjen. */
  margin: 0 calc(-1 * var(--sg-pad)) 26px !important;
  width: calc(100% + 2 * var(--sg-pad)) !important;
  max-width: none !important;
  gap: 28px !important;
  justify-content: flex-start; align-items: center !important;
}
/* Streamlit meri visinu bloka po svom sadržaju, a naši fragmenti su viši od
   te mere -- bez izričite visine reda znak i pilula ispadaju niže od ostatka. */
.st-key-sg-header [data-testid="stLayoutWrapper"],
.st-key-sg-header [data-testid="stElementContainer"],
.st-key-sg-header .stMarkdown,
.st-key-sg-header .stMarkdown > div,
.st-key-sg-header [data-testid="stMarkdownContainer"] {
  display: flex !important; align-items: center !important; height: 28px !important;
}
.st-key-sg-header .stHorizontalBlock { align-items: center !important; }
/* Streamlit vuče negativnu donju marginu na markdown-u; u redu koji centrira
   po sredini ona pomera sadržaj naniže za pola svoje vrednosti. */
.st-key-sg-header [data-testid="stMarkdownContainer"] { margin-bottom: 0 !important; }
.sg-brand, .sg-pill, .sg-avatar { line-height: 1; }
.st-key-sg-header > [data-testid="stLayoutWrapper"]:last-child { margin-left: auto; }
.st-key-sg-tools { gap: 14px !important; flex-wrap: nowrap !important; }
.st-key-sg-tools [data-testid="stLayoutWrapper"] { width: fit-content !important; }
.st-key-sg-nav { gap: 22px !important; }
.st-key-sg-brand { gap: 0 !important; }
.sg-brand { display: flex; align-items: center; gap: 10px; white-space: nowrap; }
.sg-brand span {
  font-family: var(--sg-heading); font-weight: 600; font-size: 19px; letter-spacing: 0.01em;
  line-height: 1;
}
.sg-mark { display: block; }
.st-key-sg-nav .stButton > button[data-testid="stBaseButton-tertiary"] {
  font-family: var(--sg-body) !important; font-size: 14px !important; font-weight: 400 !important;
  color: color-mix(in srgb, var(--sg-text) 75%, transparent) !important;
  padding: 0 0 3px !important; border: 0 !important; border-radius: 0 !important;
  background: transparent !important; min-height: 0 !important;
}
.st-key-sg-nav .stButton > button[data-testid="stBaseButton-tertiary"]:hover {
  color: var(--sg-accent) !important; background: transparent !important;
}

/* — pilula statusa agenta — */
.sg-pill {
  display: inline-flex; align-items: center; gap: 7px; font-size: 12px;
  padding: 4px 10px; border: 1px solid var(--sg-divider); border-radius: var(--sg-radius);
  color: color-mix(in srgb, var(--sg-text) 65%, transparent); white-space: nowrap;
}
.sg-dot { width: 5px; height: 5px; border-radius: 50%; background: var(--sg-n500); }
.sg-dot-online { background: var(--sg-accent); }
.sg-dot-warn { background: var(--sg-accent-700); }

/* — jezik i nalog — */
.st-key-sg-lang button {
  font-family: var(--sg-body) !important; font-size: 13px !important; font-weight: 400 !important;
  color: color-mix(in srgb, var(--sg-text) 75%, transparent) !important;
  border: 0 !important; background: transparent !important; padding: 2px 4px !important;
  min-height: 0 !important;
}
.st-key-sg-lang button:hover { color: var(--sg-accent) !important; }
.st-key-sg-lang [data-testid="stPopoverButton"] {
  height: 28px !important; min-height: 0 !important;
}
.st-key-sg-lang [data-testid="stPopoverButton"]::before {
  content: ""; width: 15px; height: 15px; margin-right: 6px; flex: none;
  background: currentColor;
  -webkit-mask: var(--sg-globe) center / contain no-repeat;
  mask: var(--sg-globe) center / contain no-repeat;
}
.st-key-sg-lang [data-testid="stPopoverButton"] svg { display: none !important; }
[data-testid="stPopoverBody"] {
  background: var(--sg-surface) !important; border: 1px solid var(--sg-divider) !important;
  border-radius: var(--sg-radius) !important; box-shadow: var(--sg-shadow-md) !important;
  padding: 6px !important; min-width: 186px !important;
}
[data-testid="stPopoverBody"] button,
[data-testid="stPopoverBody"] button[data-testid="stBaseButton-tertiary"] {
  font-family: var(--sg-body) !important; font-size: 13.5px !important; font-weight: 400 !important;
  color: var(--sg-text) !important; justify-content: space-between !important;
  padding: 7px 10px !important; border: 0 !important; border-radius: 2px !important;
  background: transparent !important; min-height: 0 !important;
}
[data-testid="stPopoverBody"] button:hover { background: rgba(32,31,29,0.06) !important; }
[data-testid="stPopoverBody"] button > div {
  flex: 1 1 auto; justify-content: flex-start !important; text-align: left;
}
[data-testid="stPopoverBody"] button p { line-height: 1.3 !important; }
[data-testid="stPopoverBody"] [data-testid="stVerticalBlock"] { gap: 0 !important; }
[class*="st-key-lang-"] button::after {
  font-size: 11px; letter-spacing: 0.08em;
  color: color-mix(in srgb, var(--sg-text) 50%, transparent);
}
.sg-avatar {
  width: 26px; height: 26px; border-radius: 50%; border: 1px solid var(--sg-divider);
  display: grid; place-items: center; font-family: var(--sg-heading); font-size: 12px;
}
.st-key-sg-account button {
  width: 26px !important; height: 26px !important; min-height: 26px !important;
  padding: 0 !important; border-radius: 50% !important;
  border: 1px solid var(--sg-divider) !important; background: transparent !important;
  font-family: var(--sg-heading) !important; font-size: 12px !important;
  color: var(--sg-text) !important;
}
.sg-user-name { font-size: 13px; padding: 4px 10px 8px; }

/* — sporedna kolona — */
.sg-aside {
  border-left: 1px solid var(--sg-divider); padding: 4px 0 0 26px;
  display: flex; flex-direction: column; gap: 18px; min-height: 340px;
}
.sg-aside-body {
  font-size: 12.5px; line-height: 1.6; margin: 0;
  color: color-mix(in srgb, var(--sg-text) 78%, transparent);
}
.sg-aside-dim { color: color-mix(in srgb, var(--sg-text) 70%, transparent); }
.sg-aside-foot {
  margin-top: auto; font-size: 11px;
  color: color-mix(in srgb, var(--sg-text) 45%, transparent);
}
.sg-aside-foot a { color: var(--sg-accent); text-decoration: none; }

/* — koraci i fajl — */
.sg-step { display: flex; align-items: baseline; gap: 10px; margin-bottom: 10px; }
.sg-step-no { font-family: var(--sg-heading); font-size: 13px; color: var(--sg-accent); }
.sg-file-meta {
  display: flex; align-items: center; gap: 7px; font-size: 11px; margin-top: 7px;
  color: color-mix(in srgb, var(--sg-text) 55%, transparent);
}

.st-key-sg-search [data-testid="stTextInputRootElement"] { position: relative; }
.st-key-sg-search [data-testid="stTextInputRootElement"]::before {
  content: ""; position: absolute; left: 12px; top: 50%; margin-top: -7.5px;
  width: 15px; height: 15px; opacity: 0.6; background: var(--sg-text);
  -webkit-mask: var(--sg-search) center / contain no-repeat;
  mask: var(--sg-search) center / contain no-repeat;
}
.st-key-sg-search input { padding-left: 36px !important; }

/* Streamlit markdown-u daje `margin-bottom: -15px`, čime njegova kutija ispadne
   niža nego što sadržaj jeste. U redu koji poravnava po sredini to ceo tekst
   spusti za sedam piksela -- niže od dugmeta i oznake pored njega. */
[class*="st-key-sgset-"] [data-testid="stMarkdownContainer"],
[class*="st-key-sglib-"] [data-testid="stMarkdownContainer"],
[class*="st-key-sgrow-"] [data-testid="stMarkdownContainer"],
.st-key-sg-hero-actions [data-testid="stMarkdownContainer"] {
  margin-bottom: 0 !important;
}
[class*="st-key-sgset-"] [data-testid="stMarkdownContainer"] > p,
[class*="st-key-sglib-"] [data-testid="stMarkdownContainer"] > p,
[class*="st-key-sgrow-"] [data-testid="stMarkdownContainer"] > p,
.st-key-sg-hero-actions [data-testid="stMarkdownContainer"] > p {
  margin: 0 !important;
}

/* — lista setova — */
[class*="st-key-sgset-"] {
  border-bottom: 1px solid var(--sg-divider); padding: 7px 16px !important;
}
[class*="st-key-sgset-"]:hover { background: rgba(32,31,29,0.04); }
[class*="st-key-sgset-"] button {
  font-size: 12px !important; padding: 4px 12px !important; min-height: 0 !important;
  white-space: nowrap !important;
}
[class*="st-key-sgset-"] [data-testid="stColumn"]:last-child { align-items: flex-end; }
[class*="st-key-sg-setlist-"], .st-key-sg-libtable {
  border: 1px solid var(--sg-divider) !important; border-radius: var(--sg-radius);
  padding: 0 !important; overflow-y: auto; overscroll-behavior: contain;
}
[class*="st-key-sg-setlist-"] { margin-top: 18px; }
[class*="st-key-sg-setlist-"]::-webkit-scrollbar, .st-key-sg-libtable::-webkit-scrollbar {
  width: 9px;
}
[class*="st-key-sg-setlist-"]::-webkit-scrollbar-thumb,
.st-key-sg-libtable::-webkit-scrollbar-thumb {
  background: color-mix(in srgb, var(--sg-text) 20%, transparent); border-radius: 5px;
}
.sg-set-active { color: var(--sg-accent-700); }
[class*="st-key-sgset-"]:has(.sg-set-active) {
  box-shadow: inset 2px 0 0 var(--sg-accent);
}
.sg-set-name { font-size: 14px; line-height: 1.35; padding: 2px 0; }
.sg-set-id { font-size: 11px; color: color-mix(in srgb, var(--sg-text) 55%, transparent); }
.sg-set-where { font-size: 13px; color: color-mix(in srgb, var(--sg-text) 70%, transparent); }
.sg-set-foot {
  padding: 9px 16px; font-size: 12px;
  color: color-mix(in srgb, var(--sg-text) 55%, transparent);
}
.sg-tag {
  display: inline-flex; align-items: center; font-size: 11px; letter-spacing: 0.02em;
  padding: 3px 10px; border-radius: 3px; background: var(--sg-n100); color: var(--sg-n800);
  white-space: nowrap;
}

/* — pregled pravila — */
.sg-sr-only {
  position: absolute; width: 1px; height: 1px; overflow: hidden;
  clip: rect(0 0 0 0); clip-path: inset(50%); white-space: nowrap;
}
.sg-review-head {
  display: flex; align-items: flex-end; justify-content: space-between; gap: 24px;
  border-bottom: 1px solid var(--sg-divider); padding-bottom: 16px; margin: 8px 0 20px;
}
.sg-figures { display: flex; gap: 26px; text-align: right; flex: none; }
.sg-figures .sg-figure-label { white-space: nowrap; }
.sg-figure-value { font-family: var(--sg-heading); font-size: 26px; line-height: 1; }
.sg-figure-big { font-family: var(--sg-heading); font-size: 40px; line-height: 1; font-weight: 400; }
.sg-figure-label {
  font-size: 11px; margin-top: 6px; letter-spacing: 0.1em; text-transform: uppercase;
  color: color-mix(in srgb, var(--sg-text) 55%, transparent);
}
.sg-figrow {
  display: flex; gap: 52px; padding: 18px 0; margin-bottom: 22px;
  border-top: 1px solid var(--sg-divider); border-bottom: 1px solid var(--sg-divider);
}
[class*="st-key-sgcard-"] {
  border: 1px solid var(--sg-divider); border-radius: var(--sg-radius);
  padding: 16px 18px 18px !important; margin-bottom: 20px;
}
/* Kartica se rasteže do visine reda: dve susedne završavaju u istoj liniji i
   kad jedna grupa ima devet polja a druga dva. Lanac mora biti flex do dna --
   `height: 100%` se ne razrešava kroz kolonu koja visinu dobija razvlačenjem. */
[data-testid="stColumn"]:has([class*="st-key-sgcard-"]),
[data-testid="stColumn"]:has([class*="st-key-sgcard-"]) > [data-testid="stVerticalBlock"],
[data-testid="stColumn"]:has([class*="st-key-sgcard-"]) > [data-testid="stVerticalBlock"]
  > [data-testid="stLayoutWrapper"] {
  display: flex !important; flex-direction: column !important;
}
[data-testid="stColumn"]:has([class*="st-key-sgcard-"]) > [data-testid="stVerticalBlock"],
[data-testid="stColumn"]:has([class*="st-key-sgcard-"]) > [data-testid="stVerticalBlock"]
  > [data-testid="stLayoutWrapper"],
[data-testid="stColumn"] [class*="st-key-sgcard-"] {
  flex: 1 1 auto !important;
}
.sg-card-head {
  display: flex; align-items: center; justify-content: space-between; gap: 12px;
  margin-bottom: 8px;
}
.sg-card-count { font-size: 11.5px; color: color-mix(in srgb, var(--sg-text) 55%, transparent); }
[class*="st-key-sgrow-"] { border-top: 1px solid var(--sg-divider); padding: 2px 0 !important; }
.sg-field {
  font-size: 12.5px; letter-spacing: 0.01em;
  color: color-mix(in srgb, var(--sg-text) 72%, transparent);
  overflow-wrap: anywhere;
}
.sg-field-missing { color: var(--sg-accent-700); }
.sg-colhead {
  font-size: 11px; letter-spacing: 0.08em; text-transform: uppercase;
  color: color-mix(in srgb, var(--sg-text) 55%, transparent); padding-bottom: 2px;
}
.sg-ev {
  display: flex; gap: 7px; font-size: 11.5px; line-height: 1.45;
  color: color-mix(in srgb, var(--sg-text) 60%, transparent);
}
.sg-ev-dot {
  flex: none; width: 6px; height: 6px; border-radius: 50%; margin-top: 5px;
  background: var(--sg-accent);
}
.sg-dot-medium { opacity: 0.6; }
.sg-dot-low { opacity: 0.32; }
.sg-ev-empty { color: color-mix(in srgb, var(--sg-text) 45%, transparent); font-style: normal; }
.sg-ev-page { opacity: 0.7; white-space: nowrap; }

/* — izveštaj — */
.sg-report-anchor { height: 0; scroll-margin-top: 12px; }
.sg-table { width: 100%; border-collapse: collapse; font-size: 13px; }
.sg-table th {
  text-align: left; font-size: 11px; letter-spacing: 0.08em; text-transform: uppercase;
  font-weight: 400; color: color-mix(in srgb, var(--sg-text) 60%, transparent);
  padding: 9.2px; border-bottom: 1px solid var(--sg-divider);
}
.sg-table td { padding: 9.2px; border-bottom: 1px solid var(--sg-divider); }
.sg-table tbody tr:hover { background: rgba(32,31,29,0.04); }
.sg-td-muted { font-size: 13px; color: color-mix(in srgb, var(--sg-text) 72%, transparent); }
.sg-td-right { text-align: right; }
.sg-box {
  margin-top: 24px; border: 1px solid var(--sg-divider); border-radius: var(--sg-radius);
  padding: 14px 18px 16px;
}
.sg-del-row {
  display: grid; grid-template-columns: 64px 160px 1fr; gap: 14px; font-size: 12.5px;
  padding: 6px 0; border-top: 1px solid var(--sg-divider);
}
.sg-del-index { color: color-mix(in srgb, var(--sg-text) 50%, transparent); }
.sg-del-section { color: color-mix(in srgb, var(--sg-text) 70%, transparent); }
.sg-result-name { font-size: 14px; margin-bottom: 4px; overflow-wrap: anywhere; }
.sg-result-meta {
  font-size: 11.5px; margin-bottom: 14px;
  color: color-mix(in srgb, var(--sg-text) 55%, transparent);
}
.sg-checks { display: flex; flex-direction: column; gap: 7px; margin-top: 12px; }
.sg-check { display: flex; gap: 8px; font-size: 12.5px; align-items: baseline; }

/* — biblioteka — */
.st-key-sg-import { margin-top: 10px; }
.st-key-sg-import [data-testid="stPopoverButton"] {
  font-family: var(--sg-heading) !important; font-weight: 600 !important;
  font-size: 14px !important; border: 1px solid var(--sg-divider) !important;
  border-radius: var(--sg-radius) !important; background: transparent !important;
  color: var(--sg-text) !important;
}
.st-key-sg-import [data-testid="stPopoverBody"] { min-width: 340px !important; padding: 14px !important; }
.sg-table-head {
  border: 0; margin-bottom: 2px; position: sticky; top: 0; z-index: 3;
  background: var(--sg-bg); padding-top: 8px;
}
.sg-lib-row {
  display: grid; grid-template-columns: 1.15fr 0.75fr 0.75fr 0.5fr 0.6fr 0.5fr 1.15fr;
  gap: 16px; padding: 0 8px 8px; font-size: 11px; letter-spacing: 0.08em;
  text-transform: uppercase; color: color-mix(in srgb, var(--sg-text) 60%, transparent);
  border-bottom: 1px solid var(--sg-divider);
}
[class*="st-key-sglib-"] {
  border-bottom: 1px solid var(--sg-divider); padding: 9px 8px !important;
}
.st-key-sg-libtable { margin-top: 0; }
[class*="st-key-sglib-"]:hover { background: rgba(32,31,29,0.04); }
[class*="st-key-sglibact-"] { justify-content: flex-end !important; gap: 2px !important; }
[class*="st-key-sglibact-"] button {
  font-size: 12px !important; padding: 2px 6px !important; min-height: 0 !important;
  border: 0 !important; background: transparent !important; color: var(--sg-accent) !important;
  white-space: nowrap !important;
}
[class*="st-key-sglibact-"] button:disabled { opacity: 0.45 !important; }

/* — pomoć — */
[class*="st-key-sg-helpcol-"] {
  padding: 20px 24px 8px !important; border-right: 1px solid var(--sg-divider);
  border-top: 1px solid var(--sg-divider);
}
.st-key-sg-helpcol-3 { border-right: 0; }
[class*="st-key-sg-justify-"] .stMarkdown p,
[class*="st-key-sg-justify-"] .stMarkdown li {
  font-size: 13.5px; line-height: 1.65; text-align: justify; hyphens: auto;
  color: color-mix(in srgb, var(--sg-text) 78%, transparent);
}
.sg-faq { padding: 12px 0; border-top: 1px solid var(--sg-divider); }
.sg-faq-q { font-family: var(--sg-heading); font-size: 16px; font-weight: 600; margin-bottom: 4px; }
.sg-faq-a {
  font-size: 13px; line-height: 1.6; margin: 0;
  color: color-mix(in srgb, var(--sg-text) 72%, transparent);
}

/* — kapija — */
.sg-gate { display: grid; grid-template-columns: 1fr 1fr; min-height: 420px; }
.sg-gate-lead {
  max-width: 44ch; font-size: 14px; line-height: 1.65; margin: 0;
  color: color-mix(in srgb, var(--sg-text) 78%, transparent);
}
.st-key-sg-gate-left {
  border-right: 1px solid var(--sg-divider); padding: 48px 48px 48px 0 !important;
  min-height: 420px; justify-content: center;
}
.st-key-sg-gate-right {
  padding: 48px 0 !important; max-width: 400px; min-height: 420px;
  justify-content: center; gap: 14px !important;
}
.st-key-sg-gate-left img { margin-bottom: 14px; }
.sg-or {
  display: flex; align-items: center; gap: 12px; font-size: 12px; margin: 6px 0;
  color: color-mix(in srgb, var(--sg-text) 45%, transparent);
}
.sg-or::before, .sg-or::after {
  content: ""; flex: 1; height: 1px; background: var(--sg-divider);
}
.sg-note {
  font-size: 12px; line-height: 1.6; margin: 4px 0 0;
  color: color-mix(in srgb, var(--sg-text) 60%, transparent);
}

/* ══ Streamlit-ovi widgeti u paleti sistema ══ */

/* dugmad */
.stButton > button, .stDownloadButton > button, .stFormSubmitButton > button,
.stLinkButton > a {
  font-family: var(--sg-heading) !important; font-weight: 600 !important;
  font-size: 14px !important; line-height: 1.2 !important;
  background: transparent !important; color: var(--sg-text) !important;
  border: 1px solid var(--sg-divider) !important; border-radius: var(--sg-radius) !important;
  box-shadow: none !important; padding: 9px 17px !important;
}
.stButton > button:hover, .stDownloadButton > button:hover,
.stFormSubmitButton > button:hover, .stLinkButton > a:hover {
  background: rgba(32,31,29,0.07) !important; color: var(--sg-text) !important;
  border-color: var(--sg-divider) !important;
}
button[data-testid="stBaseButton-primary"],
button[data-testid="stBaseButton-primaryFormSubmit"] {
  color: var(--sg-accent) !important; border-color: var(--sg-accent) !important;
  background: transparent !important;
}
button[data-testid="stBaseButton-primary"]:hover,
button[data-testid="stBaseButton-primaryFormSubmit"]:hover {
  background: rgba(182,130,53,0.12) !important; color: var(--sg-accent) !important;
}
button[data-testid="stBaseButton-tertiary"] {
  border-color: transparent !important; color: var(--sg-accent) !important;
  padding: 4px 6px !important;
}
button[data-testid="stBaseButton-tertiary"]:hover {
  background: rgba(182,130,53,0.10) !important; color: var(--sg-accent) !important;
}
.stButton > button:disabled, .stDownloadButton > button:disabled {
  opacity: 0.45 !important; cursor: not-allowed !important;
}

/* polja */
[data-testid="stTextInputRootElement"],
[data-testid="stNumberInputContainer"],
.react-aria-ComboBox > div[role="group"],
div[data-baseweb="input"], div[data-baseweb="select"] > div,
.stTextArea textarea {
  background: transparent !important; border: 1px solid var(--sg-divider) !important;
  border-radius: var(--sg-radius) !important; box-shadow: none !important;
  font-family: var(--sg-body) !important; font-size: 14px !important;
}
[data-testid="stTextInputRootElement"]:hover,
[data-testid="stNumberInputContainer"]:hover,
.react-aria-ComboBox > div[role="group"]:hover {
  border-color: rgba(32,31,29,0.45) !important;
}
[data-testid="stTextInputRootElement"]:focus-within,
[data-testid="stNumberInputContainer"]:focus-within,
.react-aria-ComboBox > div[role="group"]:focus-within {
  border-color: var(--sg-accent) !important;
}
.stTextInput input, .stNumberInput input, .stTextArea textarea,
.react-aria-ComboBox input, div[data-baseweb="input"] input {
  background: transparent !important; color: var(--sg-text) !important;
  font-family: var(--sg-body) !important; font-size: 14px !important;
}
.react-aria-ComboBox button { background: transparent !important; border: 0 !important; }
input::placeholder { color: color-mix(in srgb, var(--sg-text) 45%, transparent) !important; }
[data-testid="stWidgetLabel"] p {
  font-size: 12px !important; color: color-mix(in srgb, var(--sg-text) 70%, transparent) !important;
}
ul[data-testid="stSelectboxVirtualDropdown"], div[data-baseweb="popover"] ul[role="listbox"] {
  background: var(--sg-surface) !important; border: 1px solid var(--sg-divider) !important;
  border-radius: var(--sg-radius) !important; font-family: var(--sg-body) !important;
}
li[role="option"] { font-size: 13.5px !important; }
li[role="option"][aria-selected="true"], li[role="option"]:hover {
  background: rgba(182,130,53,0.12) !important;
}

/* vrednost pravila je podvučeno polje, ne kutija */
[class*="st-key-sgrow-"] [data-testid="stTextInputRootElement"],
[class*="st-key-sgrow-"] .react-aria-ComboBox > div[role="group"] {
  border: 0 !important; border-bottom: 1px solid var(--sg-divider) !important;
  border-radius: 0 !important; min-height: 28px !important;
}
[class*="st-key-sgrow-"] input {
  font-size: 13.5px !important; font-variant-numeric: tabular-nums;
  padding-left: 2px !important; padding-right: 2px !important;
}
[class*="st-key-sgrow-"] [data-testid="stTextInputRootElement"]:focus-within,
[class*="st-key-sgrow-"] .react-aria-ComboBox > div[role="group"]:focus-within {
  border-bottom-color: var(--sg-accent) !important;
}
[class*="st-key-sgrow-"] .react-aria-ComboBox button svg { height: 1rem; width: 1rem; }

/* prilaganje fajla */
section[data-testid="stFileUploaderDropzone"] {
  background: transparent !important; border: 1px solid var(--sg-divider) !important;
  border-radius: var(--sg-radius) !important; padding: 12px 16px !important;
}
section[data-testid="stFileUploaderDropzone"] button {
  font-size: 13px !important; padding: 6px 12px !important;
}
[data-testid="stFileUploaderDropzoneInstructions"] span,
[data-testid="stFileUploaderDropzoneInstructions"] small {
  font-family: var(--sg-body) !important; font-size: 12.5px !important;
  color: color-mix(in srgb, var(--sg-text) 55%, transparent) !important;
}
[data-testid="stFileUploaderFile"] { font-size: 12.5px !important; }
.st-key-sg-thesis-picked [data-testid="stFileUploaderDropzoneInstructions"] svg,
.st-key-sg-thesis-picked [data-testid="stFileUploaderDropzoneInstructions"] small {
  display: none !important;
}
.st-key-sg-thesis-picked section[data-testid="stFileUploaderDropzone"] {
  padding: 8px 12px !important; min-height: 0 !important;
}
.stFileChip {
  border: 0 !important; background: transparent !important; padding: 2px 0 !important;
}
[data-testid="stFileChipName"] {
  font-family: var(--sg-body) !important; font-size: 14px !important;
}
[data-testid="stFileChipName"] + div { display: none !important; }
.stFileChip > div:first-child {
  background: transparent !important; border-radius: 0 !important;
  width: auto !important; height: auto !important; padding: 0 !important;
}
.stFileChip > div:first-child svg {
  width: 18px !important; height: 18px !important;
  color: color-mix(in srgb, var(--sg-text) 60%, transparent) !important;
}

/* izbor izvora pravila -- segmentna kontrola */
.st-key-sg-seg [data-testid="stButtonGroup"] [role="radiogroup"] {
  display: inline-flex !important; gap: 0 !important; flex-wrap: nowrap !important;
  border: 1px solid var(--sg-divider); border-radius: var(--sg-radius); overflow: hidden;
}
.st-key-sg-seg button[data-variant="pills"] {
  font-family: var(--sg-body) !important; font-weight: 400 !important; font-size: 13px !important;
  background: transparent !important; border: 0 !important; border-radius: 0 !important;
  color: var(--sg-text) !important; padding: 7px 12px !important; min-height: 0 !important;
  box-shadow: none !important;
}
.st-key-sg-seg button[data-variant="pills"] + button[data-variant="pills"] {
  border-left: 1px solid var(--sg-divider) !important;
}
.st-key-sg-seg button[data-variant="pills"]:hover {
  background: rgba(32,31,29,0.07) !important;
}
.st-key-sg-seg button[data-variant="pills"][aria-checked="true"],
.st-key-sg-seg button[data-variant="pills"][data-selected="true"] {
  color: var(--sg-accent) !important;
  box-shadow: inset 0 0 0 1px var(--sg-accent) !important;
  background: transparent !important;
}

/* potvrde -- tačka, ne kvadratić */
.stCheckbox label > div:not([data-testid="stWidgetLabel"]) {
  border-radius: 50% !important; width: 16px !important; height: 16px !important;
  min-width: 16px !important; border: 1.5px solid var(--sg-divider) !important;
  background: transparent !important; box-shadow: none !important;
}
.stCheckbox label:hover > div:not([data-testid="stWidgetLabel"]) {
  border-color: var(--sg-accent) !important;
}
.stCheckbox label[data-selected="true"] > div:not([data-testid="stWidgetLabel"]) {
  background: var(--sg-accent) !important; border-color: var(--sg-accent) !important;
  box-shadow: inset 0 0 0 4px var(--sg-bg) !important;
}
.stCheckbox label > div:not([data-testid="stWidgetLabel"]) svg { display: none !important; }
.stCheckbox label p { font-size: 14px !important; color: var(--sg-text) !important; }

/* harmonika, poruke, status */
[data-testid="stExpander"] details {
  border: 1px solid var(--sg-divider) !important; border-radius: var(--sg-radius) !important;
  background: transparent !important;
}
[data-testid="stExpander"] summary p {
  font-family: var(--sg-heading) !important; font-weight: 600 !important; font-size: 15px !important;
}
[data-testid="stAlert"], [data-testid="stAlertContainer"],
[data-testid="stNotification"], [data-testid="stNotificationContentInfo"],
[data-testid="stNotificationContentWarning"], [data-testid="stNotificationContentError"],
[data-testid="stNotificationContentSuccess"] {
  border-radius: var(--sg-radius) !important;
  background: transparent !important; color: var(--sg-text) !important;
  font-size: 13px !important; box-shadow: none !important;
}
[data-testid="stAlertContainer"] { border: 1px solid var(--sg-divider) !important; }
[data-testid="stAlert"] p, [data-testid="stAlertContainer"] p { font-size: 13px !important; }
[data-testid="stAlertContainer"] svg { color: var(--sg-accent) !important; }
[data-testid="stCaptionContainer"] p, .stCaption p {
  font-size: 12px !important; color: color-mix(in srgb, var(--sg-text) 55%, transparent) !important;
}
.stMarkdown code {
  font-size: 12.5px; background: var(--sg-n100); color: var(--sg-n800);
  border-radius: 3px; padding: 1px 5px;
}
.stMarkdown a, .sg-aside a { color: var(--sg-accent); }

/* akcije */
.st-key-sg-apply, .st-key-sg-hero-actions {
  gap: 14px !important; flex-wrap: wrap !important; align-items: center !important;
}
/* Lepljivo dno mora da stoji na omotaču koji Streamlit pravi oko kontejnera:
   sam kontejner je jedino dete tog omotača, pa mu je blok toliko visok koliko
   i on -- nema u čemu da se lepi. */
[data-testid="stLayoutWrapper"]:has(> .st-key-sg-apply) {
  position: sticky; bottom: 0; z-index: 20;
  /* Negativna margina traku samo pomeri ulevo; širina mora izričito da pokrije
     oba prepusta, a `flex: none` da je flex red ne stisne nazad na 100%. */
  margin: 20px calc(-1 * var(--sg-pad)) 0 !important;
  width: calc(100% + 2 * var(--sg-pad)) !important;
  max-width: none !important;
  flex: 0 0 auto !important;
  padding: 12px var(--sg-pad) !important;
  background: var(--sg-bg);
  border-top: 1px solid var(--sg-divider);
  box-shadow: 0 -3px 10px rgba(45,43,43,0.08);
}
.st-key-sg-apply div[data-testid="stSelectbox"] { min-width: 190px; }
.st-key-sg-download button { width: 100% !important; }

/* iframe za pozadinski js */
iframe[title="streamlit.components.v1.html"], iframe[height="0"] {
  display: none !important; height: 0 !important; width: 0 !important;
  position: absolute !important; border: none !important;
}

/* ══ Uži ekrani ══
   Streamlit-ove kolone se ne slažu same; svaki red koji treba da se prelomi
   dobija ključ i ovde se izričito spušta u jednu kolonu. Tabele koje bi u
   jednoj koloni prestale da budu tabele umesto toga klize vodoravno. */

@media (max-width: 1180px) {
  /* Sporedna kolona pri 260px prestaje da bude margina i postaje uska traka
     sa jednom rečju po redu -- ide pod sadržaj, kao podnožje. */
  /* Samo spoljni red, ne i svaka kolona u njemu: unutar sadržaja stoje redovi
     koraka i redovi liste setova, koji ne smeju da se prelome ovde. */
  .st-key-sg-hero-row > [data-testid="stHorizontalBlock"],
  .st-key-sg-hero-row > * > [data-testid="stHorizontalBlock"] { flex-wrap: wrap !important; }
  .st-key-sg-hero-row > [data-testid="stHorizontalBlock"] > [data-testid="stColumn"],
  .st-key-sg-hero-row > * > [data-testid="stHorizontalBlock"] > [data-testid="stColumn"] {
    min-width: 100% !important;
  }
  .sg-aside {
    border-left: 0; border-top: 1px solid var(--sg-divider);
    padding: 18px 0 0; min-height: 0; margin-top: 8px;
  }
  .sg-aside-foot { margin-top: 6px; }
}

@media (max-width: 1000px) {
  :root { --sg-pad: 24px; }
  .sg-h1 { font-size: 34px; }
  .sg-h2 { font-size: 26px; }
  .sg-lead, .sg-lead-narrow, .sg-lead-wide, .sg-gate-lead { max-width: none; }

  .st-key-sg-steps [data-testid="stHorizontalBlock"],
  .st-key-sg-report-row [data-testid="stHorizontalBlock"],
  .st-key-sg-lib-head [data-testid="stHorizontalBlock"],
  .st-key-sg-help-cols [data-testid="stHorizontalBlock"],
  .st-key-sg-faq [data-testid="stHorizontalBlock"],
  .st-key-sg-gate-row [data-testid="stHorizontalBlock"] { flex-wrap: wrap !important; }
  .st-key-sg-steps [data-testid="stColumn"],
  .st-key-sg-report-row [data-testid="stColumn"],
  .st-key-sg-lib-head [data-testid="stColumn"],
  .st-key-sg-help-cols [data-testid="stColumn"],
  .st-key-sg-faq [data-testid="stColumn"],
  .st-key-sg-gate-row [data-testid="stColumn"],
  [data-testid="stColumn"]:has([class*="st-key-sgcard-"]) { min-width: 100% !important; }

  [class*="st-key-sg-helpcol-"] {
    border-right: 0; border-top: 1px solid var(--sg-divider); padding: 18px 0 4px !important;
  }
  .st-key-sg-gate-left {
    border-right: 0; border-bottom: 1px solid var(--sg-divider);
    padding: 32px 0 !important; min-height: 0;
  }
  .st-key-sg-gate-right { padding: 24px 0 32px !important; min-height: 0; max-width: none; }

  /* Traka akcija u tri reda pojede pola malog ekrana; tu je bolje da ostane
     tamo gde je i da se do nje skroluje. */
  [data-testid="stLayoutWrapper"]:has(> .st-key-sg-apply) {
    position: static; box-shadow: none;
  }
  .stMainBlockContainer { padding-bottom: 48px !important; }

  /* Tabela biblioteke ima sedam kolona -- u jednoj koloni bi bila spisak
     bez zaglavlja, pa radije klizi. */
  .st-key-sg-libtable { overflow-x: auto; }
  .st-key-sg-libtable [data-testid="stHorizontalBlock"],
  .st-key-sg-libtable .sg-lib-row { min-width: 940px; }

  [class*="st-key-sg-setlist-"] { overflow-x: auto; }
  [class*="st-key-sgset-"] [data-testid="stHorizontalBlock"] { min-width: 560px; }
}

@media (max-width: 760px) {
  :root { --sg-pad: 16px; }
  .st-key-sg-header { gap: 16px !important; }
  .st-key-sg-nav { gap: 16px !important; }
  .st-key-sg-tools { gap: 10px !important; }
  .sg-figrow { gap: 28px; flex-wrap: wrap; }
  .sg-review-head { flex-wrap: wrap; }
  .sg-figures { gap: 20px; }
  .st-key-sg-seg [data-testid="stButtonGroup"] [role="radiogroup"] { flex-wrap: wrap !important; }
  .sg-del-row { grid-template-columns: 52px 1fr; }
  .sg-del-row span:nth-child(2) { display: none; }
  /* Ime polja, vrednost i citat jedno pod drugim -- tri kolone na telefonu
     ostavljaju svakoj po nekoliko znakova. */
  [class*="st-key-sgrow-"] [data-testid="stHorizontalBlock"] { flex-wrap: wrap !important; }
  [class*="st-key-sgrow-"] [data-testid="stColumn"] { min-width: 100% !important; }
  [class*="st-key-sgrow-"] { padding: 6px 0 !important; }
  /* Obostrano poravnanje na uskoj koloni pravi reke belina; ovde je levo
     poravnanje čitljivije od simetrije. */
  [class*="st-key-sg-justify-"] .stMarkdown p,
  [class*="st-key-sg-justify-"] .stMarkdown li { text-align: left; }
}
</style>
"""


def _inject_custom_css(active_page: str) -> None:
    """Jedan stil za celu aplikaciju; aktivnu stavku navigacije nosi poslednje pravilo."""
    rules = [
        f".st-key-sg-nav .st-key-nav-{active_page} .stButton > "
        'button[data-testid="stBaseButton-tertiary"] {'
        "color: var(--sg-accent) !important;"
        "border-bottom: 1px solid var(--sg-accent) !important;"
        "border-radius: 0 !important;}"
    ]
    # Kôd jezika stoji desno u redu menija; sadržaj se ne može zakačiti na
    # dugme kao atribut, pa se pravilo ispisuje po jeziku.
    rules += [
        f'.st-key-lang-{code} button::after {{content: "{code.upper()}";}}'
        for code in i18n.available_languages()
    ]
    st.markdown(_CSS + "<style>" + "".join(rules) + "</style>", unsafe_allow_html=True)


# --------------------------------------------------------------------------


def main() -> None:
    _apply_language()

    page = st.session_state.get("nav_page") or "format"
    if page not in NAV_PAGES:
        page = "format"
    _inject_custom_css(page)

    if not identity.require_gate(st):
        return

    user = identity.current_user(st)
    _app_header(user, page)

    if page == "help":
        page_help()
    elif page == "library":
        page_library(user)
    else:
        page_format(user)


if __name__ == "__main__":
    main()
