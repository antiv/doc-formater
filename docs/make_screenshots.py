#!/usr/bin/env python3
"""Regenerate the screenshots used in the README.

    .venv/bin/pip install playwright && .venv/bin/playwright install chromium
    .venv/bin/python docs/make_screenshots.py

Everything shown is generated here: a synthetic thesis, a synthetic style guide
and a throwaway rule library in a temp directory. No real document ever reaches
the images — the working tree during development holds other people's theses.

The app is started with its own `RULES_LIBRARY_DIR` so the developer's real
library is neither shown nor modified.
"""

from __future__ import annotations

import io
import os
import shutil
import socket
import subprocess
import sys
import tempfile
import time
import zlib
from datetime import datetime, timedelta, timezone
from pathlib import Path

import docx
from docx.shared import Cm

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from styleguard.library import RuleLibrary  # noqa: E402
from styleguard.rules import Institution, load_rule_set  # noqa: E402

OUT_DIR = Path(__file__).resolve().parent / "images"
PORT = 8555
VIEWPORT = {"width": 1500, "height": 1000}


# --------------------------------------------------------------------------
# Synthetic documents
# --------------------------------------------------------------------------


def _png_1px() -> bytes:
    def chunk(tag: bytes, payload: bytes) -> bytes:
        body = tag + payload
        return (
            len(payload).to_bytes(4, "big")
            + body
            + (zlib.crc32(body) & 0xFFFFFFFF).to_bytes(4, "big")
        )

    ihdr = chunk(b"IHDR", (8).to_bytes(4, "big") + (8).to_bytes(4, "big") + bytes([8, 2, 0, 0, 0]))
    raw = b"".join(b"\x00" + b"\xc8\xd4\xe4" * 8 for _ in range(8))
    return b"\x89PNG\r\n\x1a\n" + ihdr + chunk(b"IDAT", zlib.compress(raw)) + chunk(b"IEND", b"")


def build_thesis(path: Path) -> None:
    """A small thesis with everything the formatter reacts to."""
    document = docx.Document()
    document.add_paragraph("NORTHFIELD UNIVERSITY")
    document.add_paragraph("Faculty of Design")
    document.add_paragraph("")
    document.add_paragraph("Typography in Early Printed Books")
    document.add_paragraph("A Bachelor's Thesis")
    document.add_paragraph("")

    document.add_paragraph("ACKNOWLEDGEMENTS", style="Heading 1")
    document.add_paragraph("My thanks to everyone who read the early drafts.")
    document.add_paragraph("")

    document.add_paragraph("ABSTRACT", style="Heading 1")
    document.add_paragraph(
        "This thesis surveys the typographic conventions of books printed before "
        "1500 and traces how those conventions were carried into modern practice."
    )
    document.add_paragraph("")

    document.add_paragraph("1 INTRODUCTION", style="Heading 1")
    for _ in range(2):
        document.add_paragraph(
            "The earliest printed books imitated the manuscripts they replaced, "
            "down to the abbreviations and the spacing between words. Only later "
            "did printers begin to treat the page as a surface with rules of its own."
        )
        document.add_paragraph("")

    document.add_paragraph("1.1 Scope of the study", style="Heading 2")
    document.add_paragraph(
        "The corpus is limited to books printed in the German-speaking lands, "
        "which keeps the sample small enough to examine page by page."
    )
    document.add_paragraph("")

    image_paragraph = document.add_paragraph()
    image_paragraph.add_run().add_picture(io.BytesIO(_png_1px()), width=Cm(4))
    document.add_paragraph("Figure 1: A page from a 1476 edition")
    document.add_paragraph("Source: Author's photograph")
    document.add_paragraph("")

    document.add_paragraph("2 FINDINGS", style="Heading 1")
    document.add_paragraph("Table 1: Measured line spacing by decade")
    table = document.add_table(rows=3, cols=3)
    for column, header in enumerate(["Decade", "Books", "Mean spacing"]):
        table.cell(0, column).text = header
    for row, values in enumerate([["1470s", "12", "1.18"], ["1480s", "27", "1.22"]], start=1):
        for column, value in enumerate(values):
            table.cell(row, column).text = value
    document.add_paragraph("")
    document.add_paragraph(
        "Line spacing tightened as printers gained confidence in the medium, "
        "though the change is far from uniform across workshops."
    )
    document.add_paragraph("")

    document.add_paragraph("REFERENCES", style="Heading 1")
    for entry in [
        "Barrow, L. (1998). The Printed Page. Ashgate.",
        "Ekstrom, J. and Vance, R. (2004). Early Type Design. Harbour Press.",
        "Nadel, P. (2011). Reading the Incunabula. University Books.",
    ]:
        document.add_paragraph(entry)
    document.add_paragraph("")

    document.add_paragraph("APPENDICES", style="Heading 1")
    document.add_paragraph("")
    document.add_paragraph("1. Which workshops were sampled?")
    document.add_paragraph("2. How were the measurements taken?")

    document.save(str(path))


def build_style_guide(path: Path) -> None:
    """A style guide stating rules in prose and in a table."""
    document = docx.Document()
    document.add_paragraph("NORTHFIELD UNIVERSITY")
    document.add_paragraph("Faculty of Design")
    document.add_paragraph("THESIS FORMATTING GUIDELINES")
    document.add_paragraph("")

    document.add_paragraph("1 GENERAL", style="Heading 1")
    document.add_paragraph(
        "The body text is set in Times New Roman at 12 pt. Line spacing is 1.15. "
        "Paragraphs are separated by 12 pt of space after; there is no space before."
    )
    document.add_paragraph(
        "Body text must be justified. Italics are not permitted in the body text, "
        "and there must be no empty lines between paragraphs."
    )
    document.add_paragraph("")

    document.add_paragraph("2 PAGE SETUP", style="Heading 1")
    document.add_paragraph(
        "Paper size is A4. Margins are: top 2.5 cm, bottom 2.5 cm, inside 3 cm, "
        "outside 2.5 cm. Mirror margins are used. The first page carries no header."
    )
    document.add_paragraph("")

    document.add_paragraph("3 HEADINGS", style="Heading 1")
    table = document.add_table(rows=4, cols=4)
    for column, header in enumerate(["Level", "Size", "Style", "New page"]):
        table.cell(0, column).text = header
    for row, values in enumerate(
        [
            ["Level 1", "14 pt", "bold, capitals", "yes"],
            ["Level 2", "12 pt", "bold", "no"],
            ["Level 3 and below", "12 pt", "regular", "no"],
        ],
        start=1,
    ):
        for column, value in enumerate(values):
            table.cell(row, column).text = value
    document.add_paragraph("")

    document.add_paragraph("4 FIGURES AND TABLES", style="Heading 1")
    document.add_paragraph(
        "A figure caption is placed below the figure, centred, bold, 12 pt. "
        "A table caption is placed above the table, centred, bold, 12 pt."
    )
    document.add_paragraph(
        "The source is given below the figure. The source line is centred and set "
        "in regular type at 12 pt."
    )
    document.add_paragraph("")

    document.add_paragraph("5 REFERENCES", style="Heading 1")
    document.add_paragraph(
        "The chapter is titled REFERENCES and starts on a new page. Entries are "
        "justified, 12 pt, line spacing 1.15, with no space before or after."
    )
    document.add_paragraph("")

    document.add_paragraph("6 STRUCTURE", style="Heading 1")
    document.add_paragraph(
        "A thesis contains: ACKNOWLEDGEMENTS, ABSTRACT, INTRODUCTION, the main "
        "chapters, REFERENCES and APPENDICES."
    )
    document.save(str(path))


def seed_library(directory: Path) -> None:
    """A few sets so the library page shows something worth looking at."""
    library = RuleLibrary(directory)
    base = load_rule_set(ROOT / "presets" / "ameu.json")

    demos = [
        ("Northfield University — Faculty of Design (bachelor's thesis)",
         Institution(university="Northfield University", faculty="Faculty of Design",
                     document_type="bachelor's thesis", language="en"),
         "ana@example.com", "Ana Whitfield", 0),
        ("Northfield University — Faculty of Design (master's thesis)",
         Institution(university="Northfield University", faculty="Faculty of Design",
                     document_type="master's thesis", language="en"),
         "ana@example.com", "Ana Whitfield", 3),
        ("Riverbend Institute — School of Architecture",
         Institution(university="Riverbend Institute", faculty="School of Architecture",
                     document_type="diploma thesis", language="en"),
         "tomas@example.com", "Tomáš Beran", 11),
    ]

    for name, institution, owner, owner_name, days_ago in demos:
        rule_set = base.model_copy(deep=True)
        rule_set.meta.id = library.unique_id(name)
        rule_set.meta.display_name = name
        rule_set.meta.institution = institution
        rule_set.meta.owner = owner
        rule_set.meta.owner_name = owner_name
        rule_set.meta.origin = "extracted"
        library.save(rule_set)
        # `save` stamps `updated_at`; rewrite it so the table shows a spread.
        rule_set.meta.updated_at = datetime.now(timezone.utc) - timedelta(days=days_ago)
        from styleguard.rules import dump_rule_set

        dump_rule_set(rule_set, library.path_for(rule_set.meta.id))


# --------------------------------------------------------------------------
# Driving the app
# --------------------------------------------------------------------------


def wait_for_port(port: int, timeout: float = 90.0) -> None:
    deadline = time.time() + timeout
    while time.time() < deadline:
        with socket.socket() as probe:
            if probe.connect_ex(("127.0.0.1", port)) == 0:
                return
        time.sleep(1)
    raise RuntimeError(f"Streamlit did not come up on port {port}")


def scroll_to(page, locator, offset: int = 90) -> None:
    """Put `locator` near the top of the viewport, not merely inside it."""
    locator.evaluate("element => element.scrollIntoView({block: 'start'})")
    page.mouse.wheel(0, -offset)
    page.wait_for_timeout(900)


def trim_blank_bottom(path: Path, padding: int = 32) -> None:
    """Crop trailing empty page below the content.

    Streamlit pages end in a tall stretch of background; left in, it makes every
    screenshot look mostly empty in the README.
    """
    from PIL import Image

    with Image.open(path) as image:
        rgb = image.convert("RGB")
        width, height = rgb.size
        # The page's background, sampled clear of the marginalia column.
        background = rgb.getpixel((int(width * 0.75), height - 4))
        last_content = 0
        for y in range(height - 1, -1, -1):
            row = [rgb.getpixel((x, y)) for x in range(int(width * 0.30), width, 24)]
            if any(pixel != background for pixel in row):
                last_content = y
                break
        cut = min(height, last_content + padding)
        if cut < height - padding:
            rgb.crop((0, 0, width, cut)).save(path)


def shoot(page, name: str, full_page: bool = True) -> None:
    path = OUT_DIR / name
    page.wait_for_timeout(1200)
    page.screenshot(path=str(path), full_page=full_page)
    trim_blank_bottom(path)
    print(f"  {path.relative_to(ROOT)}")


def main() -> int:
    from playwright.sync_api import sync_playwright

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    workspace = Path(tempfile.mkdtemp(prefix="styleguard-shots-"))
    thesis = workspace / "thesis.docx"
    guide = workspace / "formatting-guidelines.docx"
    library_dir = workspace / "library"

    build_thesis(thesis)
    build_style_guide(guide)
    seed_library(library_dir)
    print(f"Demo material in {workspace}")

    env = {
        **os.environ,
        "RULES_LIBRARY_DIR": str(library_dir),
        "STREAMLIT_BROWSER_GATHER_USAGE_STATS": "false",
    }
    # APP_PASSWORD/ADMIN_EMAILS would put a gate in front of the screenshots.
    # MATE_PAT is deliberately kept: with an agent configured the evidence
    # column shows verbatim quotes, which is the feature worth showing.
    for key in ("APP_PASSWORD", "ADMIN_EMAILS"):
        env.pop(key, None)

    server = subprocess.Popen(
        [sys.executable, "-m", "streamlit", "run", "app.py",
         f"--server.port={PORT}", "--server.headless=true"],
        cwd=str(ROOT), env=env,
        stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
    )

    try:
        wait_for_port(PORT)
        with sync_playwright() as playwright:
            browser = playwright.chromium.launch()
            page = browser.new_context(
                viewport=VIEWPORT, device_scale_factor=2,
                extra_http_headers={"Accept-Language": "en-GB,en;q=0.9"},
            ).new_page()
            page.goto(f"http://127.0.0.1:{PORT}", wait_until="networkidle")
            page.wait_for_selector("text=Format a document", timeout=60_000)

            # 1 — landing page with the two uploads
            page.set_input_files("input[type=file]", str(thesis))
            page.wait_for_timeout(2500)
            shoot(page, "01-format-page.png")

            # 2 — rule extraction from the style guide
            page.get_by_text("Upload a style guide", exact=True).click()
            page.wait_for_timeout(1500)
            uploads = page.locator("input[type=file]")
            uploads.nth(uploads.count() - 1).set_input_files(str(guide))
            page.wait_for_timeout(2500)
            page.get_by_role("button", name="Extract rules").click()
            # The rules review opens as cards, one per group, already expanded --
            # the evidence column is on screen without a click.
            page.wait_for_selector(".sg-review-head", timeout=300_000)
            page.wait_for_timeout(2500)
            scroll_to(page, page.locator(".sg-review-head").first)
            shoot(page, "02-rules-review.png", full_page=False)

            # 3 — the report, after formatting
            page.get_by_role("button", name="Format", exact=True).click()
            page.wait_for_selector("h2:has-text('What was changed')", timeout=300_000)
            page.wait_for_timeout(2500)
            scroll_to(page, page.locator("h2:has-text('What was changed')").first)
            shoot(page, "03-report.png", full_page=False)

            # 4 — the library, with owners
            page.get_by_text("Rule library", exact=True).first.click()
            # The dataframe grid renders into a canvas, so its header cells never
            # become "visible" to Playwright; wait on the page heading instead.
            page.wait_for_selector("h2:has-text('Rule library')", timeout=60_000)
            page.wait_for_timeout(2500)
            shoot(page, "04-library.png")

            browser.close()
    finally:
        server.terminate()
        try:
            server.wait(timeout=15)
        except subprocess.TimeoutExpired:
            server.kill()
        shutil.rmtree(workspace, ignore_errors=True)

    print("Done.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
